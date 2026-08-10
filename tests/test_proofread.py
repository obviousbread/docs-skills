"""Behavior checks for the portable proofread skill."""

import importlib.util
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPT = REPO_ROOT / "skills" / "proofread" / "scripts" / "redline_docx.py"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _load_proofread():
    spec = importlib.util.spec_from_file_location("proofread_redline", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _rewrite_document_xml(path, update):
    with zipfile.ZipFile(path) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}
    root = etree.fromstring(parts["word/document.xml"])
    update(root)
    parts["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in parts.items():
            archive.writestr(name, payload)


def _write_edits(path, edits):
    path.write_text(
        json.dumps({"author": "Тестовый корректор", "edits": edits}, ensure_ascii=False),
        encoding="utf-8",
    )


def test_extract_marks_soft_hyphens_visibly(tmp_path, capsys):
    proofread = _load_proofread()
    src = tmp_path / "soft-hyphens.docx"
    doc = Document()
    doc.add_paragraph("намеренныйперенос")
    doc.add_paragraph("встроенный\u00adперенос")
    doc.save(src)

    def add_ooxml_soft_hyphen(root):
        text = next(node for node in root.iter(W + "t") if node.text == "намеренныйперенос")
        run = text.getparent()
        text.text = "намеренный"
        soft_hyphen = etree.Element(W + "softHyphen")
        suffix = etree.Element(W + "t")
        suffix.text = "перенос"
        run.insert(run.index(text) + 1, soft_hyphen)
        run.insert(run.index(soft_hyphen) + 1, suffix)

    _rewrite_document_xml(src, add_ooxml_soft_hyphen)

    assert proofread.cmd_extract(src) is None
    output = capsys.readouterr().out
    assert "намеренный⟨SHY⟩перенос" in output
    assert "встроенный⟨SHY⟩перенос" in output


def test_review_comments_only_non_obvious_edits(tmp_path):
    proofread = _load_proofread()
    src = tmp_path / "input.docx"
    dst = tmp_path / "output.docx"
    edits_path = tmp_path / "edits.json"
    doc = Document()
    doc.add_paragraph("Очевидная опечаткка.")
    doc.add_paragraph("Неоднозначный вариант формулировки.")
    doc.save(src)
    _write_edits(
        edits_path,
        [
            {"block": 1, "old": "опечаткка", "new": "опечатка"},
            {
                "block": 2,
                "old": "вариант",
                "new": "смысл",
                "comment": "Выбор уточняет неоднозначное прочтение.",
            },
        ],
    )

    assert proofread.cmd_review(src, dst, edits_path) == 0
    assert proofread.cmd_audit(dst, "Тестовый корректор") == 0

    with zipfile.ZipFile(dst) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
        settings = etree.fromstring(archive.read("word/settings.xml"))
        comments = etree.fromstring(archive.read("word/comments.xml"))
    ns = {"w": W[1:-1]}
    revisions = root.xpath(".//w:ins//w:r | .//w:del//w:r", namespaces=ns)
    yellow = root.xpath(
        ".//w:ins//w:r[w:rPr/w:highlight[@w:val='yellow']] | "
        ".//w:del//w:r[w:rPr/w:highlight[@w:val='yellow']]",
        namespaces=ns,
    )
    assert settings.find(W + "trackChanges") is not None
    assert len(revisions) == len(yellow) > 0
    assert len(comments.findall(W + "comment")) == 1
    assert len(root.findall(".//" + W + "commentRangeStart")) == 1


def test_review_without_comments_keeps_only_tracked_changes(tmp_path):
    proofread = _load_proofread()
    src = tmp_path / "input.docx"
    dst = tmp_path / "output.docx"
    edits_path = tmp_path / "edits.json"
    doc = Document()
    doc.add_paragraph("Лишняяя буква.")
    doc.save(src)
    _write_edits(edits_path, [{"block": 1, "old": "Лишняяя", "new": "Лишняя"}])

    assert proofread.cmd_review(src, dst, edits_path) == 0
    assert proofread.cmd_audit(dst, "Тестовый корректор") == 0
    with zipfile.ZipFile(dst) as archive:
        assert "word/comments.xml" not in archive.namelist()
        settings = etree.fromstring(archive.read("word/settings.xml"))
    assert settings.find(W + "trackChanges") is not None
    assert proofread._default_review_dst(src).endswith(" (с правками).docx")
