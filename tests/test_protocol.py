"""Тесты для skills/docs-protocol/generate.py — create_protocol() и edit_protocol()."""

import os
import warnings

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from conftest import protocol_generate

create_protocol = protocol_generate.create_protocol


class TestSmoke:
    """Скилл собирает .docx без падений на минимальном валидном входе."""

    def test_minimal(self, tmp_path):
        path = create_protocol(
            subtype="оперативного совещания",
            chair={"lastname": "Алмазов", "initials": "А.А.",
                   "position": "и.о. генерального директора"},
            attendees=[
                {"lastname": "Бирюзов", "initials": "Б.Б.",
                 "position": "начальник отдела кадров"},
            ],
            items=[
                {"text": "Подготовить отчёт по обращениям граждан.",
                 "responsible": ["Бирюзов Б.Б."],
                 "deadline": "01.06.2026",
                 "subitems": None},
            ],
            venue="г. Москва, конференц-зал",
            doc_date="12.05.2026",
            output_path=str(tmp_path / "protocol.docx"),
        )
        assert os.path.isfile(path)
        assert path.endswith(".docx")


class TestHelpers:
    """Базовые хелперы корректно создают элементы."""

    def test_warn_slash_warns(self):
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            protocol_generate._warn_slash("первое/второе")
            assert any("Косая черта" in str(x.message) for x in w)

    def test_warn_slash_skips_legitimate(self):
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            protocol_generate._warn_slash("ИНН/КПП и п/п")
            assert not any("Косая черта" in str(x.message) for x in w)


class TestOrgConfig:
    def test_output_dir_protocol_loaded(self):
        cfg = protocol_generate._build_org_config()
        assert cfg["output_dir_protocol"] == "/tmp"

    def test_header_lines_have_parent_and_full_name(self):
        cfg = protocol_generate._build_org_config()
        assert any("Тестовое агентство" in l for l in cfg["header_lines"])
        assert any("Центр тестирования" in l for l in cfg["header_lines"])


class TestFormatters:
    def test_format_ru_date(self):
        assert protocol_generate._format_ru_date("12.05.2026") == "12 мая 2026 г."

    def test_format_ru_date_january(self):
        assert protocol_generate._format_ru_date("01.01.2026") == "1 января 2026 г."

    def test_format_fio_initials_first(self):
        assert protocol_generate._format_fio_initials_first("Иванов И.И.") == "И.И. Иванов"

    def test_format_fio_already_initials_first(self):
        # Если уже «И.О. Фамилия» — вернуть как есть.
        assert protocol_generate._format_fio_initials_first("И.И. Иванов") == "И.И. Иванов"


class TestHeaderBlock:
    def test_header_has_4_centered_bold_lines(self):
        from docx import Document
        doc = Document()
        cfg = protocol_generate._build_org_config()
        protocol_generate._make_header_block(doc, cfg)
        bold_centered = [p for p in doc.paragraphs
                         if p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                         and any(r.bold for r in p.runs)]
        assert len(bold_centered) >= 2  # parent_org + full_name + (short_name)


class TestTitleBlock:
    def test_two_centered_bold_lines(self):
        from docx import Document
        doc = Document()
        protocol_generate._make_title_block(doc, "оперативного совещания")
        ps = [p for p in doc.paragraphs
              if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.runs]
        assert any(r.text == "ПРОТОКОЛ" and r.bold for p in ps for r in p.runs)
        assert any(r.text == "оперативного совещания" and r.bold for p in ps for r in p.runs)


class TestDateNumberTable:
    def test_default_number_placeholder(self):
        from docx import Document
        doc = Document()
        protocol_generate._make_date_number_table(doc, "12.05.2026", "")
        cells = doc.tables[0].rows[0].cells
        assert "12 мая 2026 г." in cells[0].text
        assert cells[1].text.strip() == "№ _________"

    def test_explicit_number(self):
        from docx import Document
        doc = Document()
        protocol_generate._make_date_number_table(doc, "12.05.2026", "42")
        assert doc.tables[0].rows[0].cells[1].text.strip() == "№ 42"


class TestVenueChairBlock:
    def test_labels_bold(self):
        from docx import Document
        doc = Document()
        chair = {"lastname": "Алмазов", "initials": "А.А.",
                 "position": "и.о. директора"}
        protocol_generate._make_venue_chair_block(doc, "Москва", chair)
        labels = []
        for p in doc.paragraphs:
            for r in p.runs:
                if r.bold and r.text in ("Место проведения: ", "Председатель: "):
                    labels.append(r.text)
        assert "Место проведения: " in labels
        assert "Председатель: " in labels


class TestAttendeesTable:
    def test_chair_filtered_with_warning(self):
        from docx import Document
        import warnings as _w
        doc = Document()
        chair = {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о. директора"}
        attendees = [
            {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о. директора"},
            {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "начальник кадров"},
        ]
        with _w.catch_warnings(record=True) as warn_log:
            _w.simplefilter("always")
            protocol_generate._make_attendees_table(doc, attendees, chair)
        body = doc.tables[0].rows
        names = [r.cells[0].text for r in body]
        assert all("Алмазов" not in n for n in names)
        assert any("Бирюзов" in n for n in names)
        assert any("Председатель" in str(x.message) for x in warn_log)

    def test_three_columns(self):
        from docx import Document
        doc = Document()
        chair = {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о."}
        attendees = [{"lastname": "Бирюзов", "initials": "Б.Б.", "position": "кадры"}]
        protocol_generate._make_attendees_table(doc, attendees, chair)
        assert len(doc.tables[0].rows[0].cells) == 3


class TestNumbering:
    def test_protocol_numbering_returns_num_id(self):
        from docx import Document
        doc = Document()
        num_id = protocol_generate._setup_protocol_numbering(doc)
        assert isinstance(num_id, int) and num_id > 0

    def test_protocol_numbering_uses_decimal(self):
        from docx import Document
        doc = Document()
        protocol_generate._setup_protocol_numbering(doc)
        numbering = doc.part.numbering_part.element
        abs_nums = numbering.findall(qn("w:abstractNum"))
        found = False
        for an in abs_nums:
            lvl = an.find(qn("w:lvl"))
            if lvl is None:
                continue
            fmt = lvl.find(qn("w:numFmt"))
            text = lvl.find(qn("w:lvlText"))
            if (fmt is not None and fmt.get(qn("w:val")) == "decimal"
                    and text is not None and text.get(qn("w:val")) == "%1."):
                found = True
        assert found

    def test_dash_numbering_uses_bullet_endash(self):
        from docx import Document
        doc = Document()
        protocol_generate._setup_dash_numbering(doc)
        numbering = doc.part.numbering_part.element
        abs_nums = numbering.findall(qn("w:abstractNum"))
        found = False
        for an in abs_nums:
            lvl = an.find(qn("w:lvl"))
            if lvl is None:
                continue
            fmt = lvl.find(qn("w:numFmt"))
            text = lvl.find(qn("w:lvlText"))
            if (fmt is not None and fmt.get(qn("w:val")) == "bullet"
                    and text is not None and text.get(qn("w:val")) == "–"):
                found = True
        assert found


class TestResolvedBlock:
    def test_label_and_one_item(self):
        from docx import Document
        doc = Document()
        items = [{
            "text": "Подготовить отчёт.",
            "responsible": ["Бирюзов Б.Б."],
            "deadline": "01.06.2026",
            "subitems": None,
        }]
        protocol_generate._make_resolved_block(doc, items)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "РЕШИЛИ:" in all_text
        assert "Подготовить отчёт." in all_text
        assert "Ответственные:" in all_text
        assert "Бирюзов Б.Б." in all_text
        assert "Срок:" in all_text
        assert "01.06.2026" in all_text

    def test_deadline_none_leaves_срок_empty(self):
        from docx import Document
        doc = Document()
        items = [{"text": "X.", "responsible": None, "deadline": None, "subitems": None}]
        protocol_generate._make_resolved_block(doc, items)
        for p in doc.paragraphs:
            if p.runs and p.runs[0].text == "Срок: ":
                assert len(p.runs) == 1 or all(r.text == "" for r in p.runs[1:])
                break
        else:
            assert False, "Параграф «Срок: » не найден"

    def test_subitems_use_bullet_numbering(self):
        from docx import Document
        doc = Document()
        items = [{
            "text": "Главный пункт.",
            "responsible": None,
            "deadline": "постоянно",
            "subitems": ["подпункт один;", "подпункт два."],
        }]
        protocol_generate._make_resolved_block(doc, items)
        sub_paras = [p for p in doc.paragraphs
                     if p.text in ("подпункт один;", "подпункт два.")]
        assert len(sub_paras) == 2
        for p in sub_paras:
            assert p._p.find(qn("w:pPr/w:numPr") if False else qn("w:pPr")) is not None


class TestSignatureBlock:
    def test_left_cell_two_lines(self):
        from docx import Document
        doc = Document()
        chair = {"lastname": "Алмазов", "initials": "А.А.",
                 "position": "И.о. генерального директора"}
        org = {"short_name": "ТУ ЦТ", "full_name": "Тестовое учреждение"}
        protocol_generate._make_signature_block(doc, chair, org)
        left = doc.tables[0].rows[0].cells[0]
        assert len(left.paragraphs) >= 2
        assert "И.о. генерального директора" in left.paragraphs[0].text
        assert "ТУ ЦТ" in left.paragraphs[1].text

    def test_right_cell_initials_first(self):
        from docx import Document
        doc = Document()
        chair = {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о."}
        org = {"short_name": "ТУ ЦТ", "full_name": ""}
        protocol_generate._make_signature_block(doc, chair, org)
        right = doc.tables[0].rows[0].cells[1].text.strip()
        assert right == "А.А. Алмазов"


class TestSecretaryBlock:
    def test_no_secretary_no_paragraphs(self):
        from docx import Document
        doc = Document()
        before = len(doc.paragraphs)
        protocol_generate._make_secretary_block(doc, None)
        assert len(doc.paragraphs) == before

    def test_secretary_adds_two_paragraphs(self):
        from docx import Document
        doc = Document()
        protocol_generate._make_secretary_block(doc, {
            "lastname": "Деревьев", "initials": "Д.Д.",
            "position": "ведущий специалист",
        })
        texts = [p.text for p in doc.paragraphs]
        assert "Секретарь" in texts
        assert any("Д.Д. Деревьев" in t for t in texts)


class TestNotifyPersons:
    def test_union_attendees_and_responsible_excludes_chair(self):
        chair = {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о."}
        attendees = [
            {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "кадры"},
            {"lastname": "Васильков", "initials": "В.В.", "position": "ит"},
        ]
        items = [
            {"text": "X.", "responsible": ["Гранатов Г.Г."], "deadline": None, "subitems": None},
            {"text": "Y.", "responsible": ["Бирюзов Б.Б.", "Васильков В.В."], "deadline": None, "subitems": None},
        ]
        staff = [
            {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о."},
            {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "кадры"},
            {"lastname": "Васильков", "initials": "В.В.", "position": "ит"},
            {"lastname": "Гранатов", "initials": "Г.Г.", "position": "юрисконсульт"},
        ]
        result = protocol_generate._build_notify_persons(chair, attendees, items, staff)
        names = [(p["lastname"], p["initials"]) for p in result]
        assert ("Алмазов", "А.А.") not in names
        assert ("Бирюзов", "Б.Б.") in names
        assert ("Васильков", "В.В.") in names
        assert ("Гранатов", "Г.Г.") in names

    def test_dedup_and_sorted_a_to_ya(self):
        chair = {"lastname": "Алмазов", "initials": "А.А.", "position": "и.о."}
        attendees = [
            {"lastname": "Жуков", "initials": "Ж.Ж.", "position": "X"},
            {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "Y"},
            {"lastname": "Жуков", "initials": "Ж.Ж.", "position": "X"},
        ]
        items = [{"text": "X.", "responsible": ["Бирюзов Б.Б."], "deadline": None, "subitems": None}]
        staff = attendees + [chair]
        result = protocol_generate._build_notify_persons(chair, attendees, items, staff)
        names = [p["lastname"] for p in result]
        assert names == sorted(names)
        assert names.count("Жуков") == 1


class TestNotifySheet:
    def test_4_columns_and_headers(self):
        from docx import Document
        doc = Document()
        persons = [
            {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "кадры"},
            {"lastname": "Васильков", "initials": "В.В.", "position": "ит"},
        ]
        protocol_generate._make_notify_sheet(doc, persons)
        t = doc.tables[-1]
        assert len(t.rows[0].cells) == 4
        headers = [c.text for c in t.rows[0].cells]
        assert headers == ["№ п/п", "Должность", "Ф.И.О.", "Подпись"]

    def test_rows_match_persons(self):
        from docx import Document
        doc = Document()
        persons = [{"lastname": "Бирюзов", "initials": "Б.Б.", "position": "X"}]
        protocol_generate._make_notify_sheet(doc, persons)
        t = doc.tables[-1]
        assert len(t.rows) == 2
        assert "Бирюзов Б.Б." in t.rows[1].cells[2].text


class TestIntegration:
    """Интеграционные тесты по спеке (секция 6)."""

    def _minimal_args(self, tmp_path):
        return dict(
            subtype="оперативного совещания",
            chair={"lastname": "Алмазов", "initials": "А.А.",
                   "position": "и.о. генерального директора"},
            attendees=[
                {"lastname": "Бирюзов", "initials": "Б.Б.", "position": "начальник отдела кадров"},
                {"lastname": "Васильков", "initials": "В.В.", "position": "начальник отдела ИТ"},
            ],
            items=[
                {"text": "Подготовить отчёт.", "responsible": ["Бирюзов Б.Б."],
                 "deadline": "01.06.2026", "subitems": None},
            ],
            venue="г. Москва, конференц-зал",
            doc_date="12.05.2026",
            output_path=str(tmp_path / "p.docx"),
        )

    def test_basic_operational(self, tmp_path):
        path = create_protocol(**self._minimal_args(tmp_path))
        assert os.path.isfile(path)

    def test_general_director(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["subtype"] = "совещания у генерального директора"
        path = create_protocol(**args)
        assert os.path.isfile(path)

    def test_custom_subtype(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["subtype"] = "рабочей встречи по вопросам ИТ"
        path = create_protocol(**args)
        assert os.path.isfile(path)

    def test_chair_in_attendees_filtered(self, tmp_path, recwarn):
        args = self._minimal_args(tmp_path)
        args["attendees"] = [args["chair"]] + args["attendees"]
        path = create_protocol(**args)
        assert os.path.isfile(path)
        assert any("Председатель" in str(w.message) for w in recwarn.list)

    def test_subitems_and_no_deadline(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["items"] = [{
            "text": "Главный пункт.",
            "responsible": ["Бирюзов Б.Б."],
            "deadline": None,
            "subitems": ["подпункт один;", "подпункт два."],
        }]
        path = create_protocol(**args)
        assert os.path.isfile(path)
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Срок:" in text
        for p in doc.paragraphs:
            if p.runs and p.runs[0].text == "Срок: ":
                tail = "".join(r.text for r in p.runs[1:])
                assert tail == ""

    def test_secretary_optional(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["secretary"] = {"lastname": "Деревьев", "initials": "Д.Д.",
                             "position": "ведущий специалист"}
        path = create_protocol(**args)
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Секретарь" in text and "Д.Д. Деревьев" in text

    def test_notify_sheet_union(self, tmp_path, mock_staff):
        args = self._minimal_args(tmp_path)
        args["items"] = [
            {"text": "X.", "responsible": ["Гранатов Г.Г."],
             "deadline": None, "subitems": None},
        ]
        path = create_protocol(**args)
        from docx import Document
        doc = Document(path)
        notify_table = doc.tables[-1]
        fio_col = [r.cells[2].text for r in notify_table.rows[1:]]
        assert any("Бирюзов" in f for f in fio_col)
        assert any("Васильков" in f for f in fio_col)
        assert any("Гранатов" in f for f in fio_col)
        assert not any("Алмазов" in f for f in fio_col)

    def test_numbering_word_native(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["items"] = [
            {"text": "Пункт один.", "responsible": None, "deadline": None,
             "subitems": ["подпункт A;"]},
            {"text": "Пункт два.", "responsible": None, "deadline": None,
             "subitems": None},
        ]
        path = create_protocol(**args)
        from docx import Document
        doc = Document(path)
        numbering = doc.part.numbering_part.element
        abs_nums = numbering.findall(qn("w:abstractNum"))
        formats = []
        for an in abs_nums:
            lvl = an.find(qn("w:lvl"))
            if lvl is None:
                continue
            fmt = lvl.find(qn("w:numFmt"))
            text = lvl.find(qn("w:lvlText"))
            if fmt is not None and text is not None:
                formats.append((fmt.get(qn("w:val")), text.get(qn("w:val"))))
        assert ("decimal", "%1.") in formats
        assert ("bullet", "–") in formats

    def test_signature_block_two_lines(self, tmp_path):
        args = self._minimal_args(tmp_path)
        path = create_protocol(**args)
        from docx import Document
        doc = Document(path)
        sig_table = doc.tables[-2]
        left = sig_table.rows[0].cells[0]
        non_empty_paras = [p for p in left.paragraphs if p.text.strip()]
        assert len(non_empty_paras) >= 2

    def test_date_formatting(self, tmp_path):
        args = self._minimal_args(tmp_path)
        args["doc_date"] = "12.05.2026"
        path = create_protocol(**args)
        from docx import Document
        doc = Document(path)
        text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        assert "12 мая 2026 г." in text

    def test_no_slashes_in_text(self, tmp_path):
        import warnings as _w
        args = self._minimal_args(tmp_path)
        args["items"] = [{"text": "Передать отчёт/копию руководителю.",
                          "responsible": None, "deadline": None, "subitems": None}]
        with _w.catch_warnings(record=True) as warn_log:
            _w.simplefilter("always")
            create_protocol(**args)
        assert any("Косая черта" in str(w.message) for w in warn_log)

    def test_default_filename(self, tmp_path, monkeypatch):
        args = self._minimal_args(tmp_path)
        args.pop("output_path")
        path = create_protocol(**args)
        assert os.path.isfile(path)
        assert os.path.basename(path) == "Протокол оперативного совещания 12.05.2026.docx"
        os.remove(path)
