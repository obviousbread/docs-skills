"""Tests for the standalone organizational-unit regulation generator."""

from pathlib import Path

import pytest
from docx import Document

from conftest import polozheniya_generate


def _kwargs(tmp_path, **overrides):
    values = {
        "unit_type": "отдел",
        "unit_name": "организационного развития",
        "parent_unit_name": "Управления организационного развития",
        "subordination": "начальника Управления организационного развития",
        "head_appointment": "начальника Управления организационного развития",
        "head_reports_to": "начальнику Управления организационного развития",
        "goals": ["Координация организационного развития."],
        "tasks": ["Подготовка предложений по совершенствованию процессов."],
        "functions": ["Анализ и актуализация организационных процессов."],
        "approvers": [{"position": "Начальник кадрового отдела", "name": "Примеров П.П."}],
        "output_path": str(tmp_path / "Положение.docx"),
    }
    values.update(overrides)
    return values


def _text(path):
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_department_regulation_is_complete(tmp_path):
    path = polozheniya_generate.create_polozhenie(**_kwargs(tmp_path))
    text = _text(path)

    assert Path(path).is_file()
    for expected in (
        "УТВЕРЖДАЮ",
        "Генеральный директор",
        "ПОЛОЖЕНИЕ",
        "1. Общие положения",
        "7. Заключительные положения",
        "СОГЛАСОВАНО",
        "Лист ознакомления",
        "Примеров П.П.",
    ):
        assert expected in text


def test_directorate_and_output_collision(tmp_path):
    kwargs = _kwargs(
        tmp_path,
        unit_type="управление",
        unit_name="организационного развития",
        parent_unit_name=None,
        child_units=["Отдел методологии", "Отдел анализа"],
    )
    first = polozheniya_generate.create_polozhenie(**kwargs)
    second = polozheniya_generate.create_polozhenie(**kwargs)

    assert first != second
    assert Path(first).is_file() and Path(second).is_file()
    assert "Отдел методологии" in _text(first)


def test_uses_configured_output_directory(tmp_path, mock_org_config):
    org = dict(mock_org_config, output_dir_polozheniya=str(tmp_path))
    path = polozheniya_generate.create_polozhenie(
        **_kwargs(tmp_path, org=org, output_path=None)
    )

    assert Path(path).parent == tmp_path


@pytest.mark.parametrize("unit_type", ["сектор", "", None])
def test_rejects_unknown_unit_type(tmp_path, unit_type):
    with pytest.raises(ValueError):
        polozheniya_generate.create_polozhenie(**_kwargs(tmp_path, unit_type=unit_type))


def test_requires_substantive_sections(tmp_path):
    with pytest.raises(ValueError):
        polozheniya_generate.create_polozhenie(**_kwargs(tmp_path, functions=[]))
