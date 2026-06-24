"""Generated DOCX invariants that smoke tests alone do not catch."""

from docx import Document

from conftest import di_generate, letter_generate, memo_generate, ord_generate, protocol_generate


def document_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _assert_clean_core_properties(path):
    cp = Document(path).core_properties
    values = [
        cp.author,
        cp.last_modified_by,
        cp.comments,
        cp.title,
        cp.subject,
        cp.keywords,
        cp.category,
    ]
    assert all("python-docx" not in (value or "") for value in values)
    assert (cp.author or "") == ""
    assert (cp.comments or "") == ""


def _sample_di_sections():
    return [
        {
            "title": "Общие положения",
            "points": ["Инструкция фиксирует инварианты генератора."],
        },
        {
            "title": "Трудовая функция",
            "points": ["Обеспечение проверки инвариантов."],
        },
        {
            "title": "Должностные обязанности",
            "points": ["Выполняет инвариантные обязанности."],
        },
        {"title": "Права", "points": ["Запрашивать материалы для проверки."]},
        {"title": "Ответственность", "points": ["Несет ответственность в установленном порядке."]},
        {"title": "Заключительные положения", "points": ["Инструкция действует бессрочно."]},
    ]


def test_all_generators_emit_readable_expected_text_and_clean_metadata(tmp_path):
    cases = [
        (
            ord_generate.create_ord(
                doc_type="prikaz",
                title="О проверке инвариантов",
                points=["Проверить инварианты генератора."],
                control_person="self",
                output_path=str(tmp_path / "ord.docx"),
            ),
            ["О проверке инвариантов", "Проверить инварианты генератора."],
        ),
        (
            letter_generate.create_letter(
                addressee_lines=["Директору", "ООО «Тест»"],
                body_paragraphs=["Направляем инвариантное письмо."],
                output_path=str(tmp_path / "letter.docx"),
            ),
            ["ООО «Тест»", "Направляем инвариантное письмо."],
        ),
        (
            memo_generate.create_sluzhebka(
                body_paragraphs=["Прошу проверить инварианты."],
                output_path=str(tmp_path / "memo.docx"),
            ),
            ["Служебная записка", "Прошу проверить инварианты."],
        ),
        (
            di_generate.create_di(
                position_full="специалиста по инвариантам",
                position_full_genitive="специалиста по инвариантам",
                sections=_sample_di_sections(),
                output_path=str(tmp_path / "di.docx"),
            ),
            ["ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ", "Выполняет инвариантные обязанности."],
        ),
        (
            protocol_generate.create_protocol(
                subtype="оперативного совещания",
                chair={"lastname": "Алмазов", "initials": "А.А.", "position": "и.о. директора"},
                attendees=[{"lastname": "Бирюзов", "initials": "Б.Б.", "position": "начальник отдела кадров"}],
                items=[
                    {
                        "text": "Подготовить инвариантный отчет.",
                        "responsible": ["Бирюзов Б.Б."],
                        "deadline": "01.06.2026",
                        "subitems": None,
                    }
                ],
                venue="г. Москва",
                doc_date="12.05.2026",
                output_path=str(tmp_path / "protocol.docx"),
            ),
            ["ПРОТОКОЛ", "Подготовить инвариантный отчет."],
        ),
    ]

    for path, expected_text in cases:
        text = document_text(path)
        for expected in expected_text:
            assert expected in text
        _assert_clean_core_properties(path)
