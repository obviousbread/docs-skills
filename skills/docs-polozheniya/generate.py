#!/usr/bin/env python3
"""Generate standalone regulations for organizational units.

Конфигурация организации загружается из ~/.docs-plugin/org_details.md.
"""

import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ORG_DETAILS_PATH = os.path.expanduser("~/.docs-plugin/org_details.md")


# Helpers

def _load_org_details():
    """Load the simple `key: value` config created by docs-init."""
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        raise FileNotFoundError(
            "~/.docs-plugin/org_details.md not found. Run docs-init to configure."
        )
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()
    return config


def new_document():
    doc = Document()
    props = doc.core_properties
    for field in ("author", "last_modified_by", "comments", "title", "subject", "keywords", "category"):
        setattr(props, field, "")
    return doc


def _resolve_output_path(path):
    """Return a collision-free output path."""
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    n = 2
    while os.path.exists(f"{base} {n}{ext}"):
        n += 1
    return f"{base} {n}{ext}"


def _default_output_path(org, unit_type, unit_name):
    output_dir = os.path.expanduser(org.get("output_dir_polozheniya") or "")
    unit_prep = "отделе" if unit_type == "отдел" else "управлении"
    safe_name = unit_name.replace("/", "_").replace("\\", "_")
    filename = f"Положение об {unit_prep} {safe_name}.docx"
    return os.path.join(output_dir, filename) if output_dir else filename

def _set_table_borders(table, color="000000", size="4", space="0", val="single"):
    """Apply visible borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove existing borders before applying the canonical set.
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_para(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Write formatted text into a table cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_para(doc, text, bold=False, size=14, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before=0, space_after=0, first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


# Standard unit rights

def _standard_rights(label, org_short):
    """Return the conservative default rights for a unit."""
    return [
        f"Запрашивать и получать в установленном порядке от структурных подразделений "
        f"{org_short} сведения, "
        f"необходимые для работы и выполнения своих функций.",
        f"Вносить предложения по совершенствованию форм и методов работы {label}.",
        f"Участвовать в совещаниях при рассмотрении вопросов, отнесенных к компетенции {label}.",
        f"Давать разъяснения и рекомендации по вопросам, относящимся к компетенции {label}.",
        f"Визировать проекты документов, подготовленные другими структурными подразделениями, "
        f"если в таких "
        f"документах затрагиваются вопросы, относящиеся к компетенции {label}.",
    ]


# Public API

def create_polozhenie(
    unit_type="отдел",
    unit_name="",
    parent_unit_name=None,
    subordination="",
    child_units=None,
    head_appointment="",
    head_reports_to="",
    goals=None,
    tasks=None,
    functions=None,
    extra_rights=None,
    approvers=None,
    governing_documents=None,
    approver_name=None,
    org=None,
    output_path=None,
):
    """
    Create a standalone organizational-unit regulation as a DOCX file.

    Parameters
    ----------
    unit_type : str
        Either "отдел" or "управление".
    unit_name : str
        Unit name without the unit-type word, for example "организационного развития".
    parent_unit_name : str | None
        Parent-unit name for a department; None for a directorate.
    subordination : str
        Text following «находится в подчинении».
    child_units : list[str] | None
        Child units for a directorate.
    head_appointment : str
        Text following «по представлению».
    head_reports_to : str
        Text following «непосредственно подчиняется» in section 5.4.
    goals : list[str]
        Goals for section 2.1.
    tasks : list[str]
        Tasks for section 2.2.
    functions : list[str]
        Functions for section 3.
    extra_rights : list[str] | None
        Rights added after the conservative defaults.
    approvers : list[dict]
        Ordered dictionaries containing position and name.
    governing_documents : list[str] | None
        Verified governing sources phrased in the instrumental case.
    approver_name : str
        Approver name; defaults to org_details.md.
    org : dict | None
        Organization config; defaults to org_details.md.
    output_path : str
        DOCX output path.
    """
    if not isinstance(unit_type, str) or not isinstance(unit_name, str):
        raise ValueError("unit_type and unit_name must be strings")
    unit_type = unit_type.lower()
    if unit_type not in {"отдел", "управление"}:
        raise ValueError("unit_type must be 'отдел' or 'управление'")
    if not unit_name.strip():
        raise ValueError("unit_name is required")
    if unit_type == "отдел" and not parent_unit_name:
        raise ValueError("parent_unit_name is required for an отдел")
    if not subordination.strip() or not head_appointment.strip() or not head_reports_to.strip():
        raise ValueError("subordination, head_appointment and head_reports_to are required")
    if not goals or not tasks or not functions:
        raise ValueError("goals, tasks and functions must not be empty")

    org = org or _load_org_details()
    required_org = ("full_name", "short_name", "leader_title", "leader_name_nom")
    missing = [key for key in required_org if not org.get(key)]
    if missing:
        raise ValueError(f"Missing organization fields: {', '.join(missing)}")

    ORG_FULL_NOM = org["full_name"]
    ORG_FULL = org.get("full_name_gen") or ORG_FULL_NOM
    ORG_SHORT = org["short_name"]
    leader_title = org["leader_title"]
    leader_title_gen = org.get("leader_title_gen") or leader_title
    leader_title_ins = org.get("leader_title_ins") or leader_title
    approver_name = approver_name or org["leader_name_nom"]

    extra_rights = extra_rights or []
    approvers = approvers or []
    child_units = child_units or []
    governing_documents = governing_documents or [
        "Конституцией Российской Федерации",
        "федеральными законами и иными нормативными правовыми актами Российской Федерации",
        f"локальными нормативными актами {ORG_SHORT}",
    ]
    for approver in approvers:
        if not isinstance(approver, dict) or not approver.get("position"):
            raise ValueError("Each approver must contain position; name may be empty")

    is_otdel = unit_type.lower() == "отдел"
    Label = "Отдел" if is_otdel else "Управление"
    label = "Отдела" if is_otdel else "Управления"  # родительный
    label_d = "Отделу" if is_otdel else "Управлению"  # дательный
    label_t = "Отделе" if is_otdel else "Управлении"  # предложный
    label_v = "Отдел" if is_otdel else "Управление"  # винительный

    doc = new_document()

    # Default body style.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.0

    # A4 page geometry.
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    indent = 1.25

    # Approval block.

    _add_para(doc, "«УТВЕРЖДАЮ»", size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_para(doc, leader_title, size=14,
              alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_para(doc, ORG_SHORT, size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    _add_para(doc, f"________________{approver_name}", size=14,
              alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_para(doc, "«_____»________________202 ___ года", size=14,
              alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    # Document title.

    _add_para(doc, "", size=14)
    _add_para(doc, "ПОЛОЖЕНИЕ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if is_otdel:
        _add_para(doc, f"об отделе {unit_name}", size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if parent_unit_name:
            _add_para(doc, parent_unit_name, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _add_para(doc, f"об Управлении {unit_name}", size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _add_para(doc, ORG_FULL_NOM, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Section 1: General provisions.

    _add_para(doc, "1. Общие положения", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    # 1.1
    if is_otdel:
        full_unit = (
            f"отдела {unit_name} {parent_unit_name} {ORG_FULL}"
        )
    else:
        full_unit = f"Управления {unit_name} {ORG_FULL}"

    _add_para(doc,
              f"1.1. Настоящее Положение регламентирует задачи, функции, права "
              f"и деятельность {full_unit} (далее \u2013 {Label}).",
              size=14, first_line_indent=indent)

    # 1.2
    if is_otdel:
        _add_para(doc,
                  f"1.2. {Label} является структурным подразделением {ORG_FULL} "
                  f"(далее \u2013 {ORG_SHORT}, Учреждение) и входит в состав "
                  f"{parent_unit_name} (далее \u2013 Управление). "
                  f"{Label} находится в подчинении {subordination} "
                  f"либо лиц, исполняющих их обязанности.",
                  size=14, first_line_indent=indent)
    else:
        _add_para(doc,
                  f"1.2. {Label} является структурным подразделением {ORG_FULL} "
                  f"(далее \u2013 {ORG_SHORT}, Учреждение) и находится в подчинении "
                  f"{subordination} либо лиц, исполняющих их обязанности.",
                  size=14, first_line_indent=indent)

    # 1.3
    _add_para(doc,
              f"1.3. {Label} создается, реорганизуется и ликвидируется приказом "
              f"{leader_title_gen} {ORG_SHORT} либо лица, исполняющего его обязанности.",
              size=14, first_line_indent=indent)

    next_num = 4

    # Directorate composition.
    if not is_otdel and child_units:
        units_text = ", ".join(child_units)
        _add_para(doc,
                  f"1.{next_num}. В состав {label} входит {units_text}.",
                  size=14, first_line_indent=indent)
        next_num += 1

    # Structure and staffing.
    if is_otdel:
        _add_para(doc,
                  f"1.{next_num}. Структура {label} и штатная численность работников {label} "
                  f"утверждаются {leader_title_ins} {ORG_SHORT} по предложению "
                  f"{head_appointment}, исходя из возложенных на {Label} функций и задач "
                  f"с учетом рекомендуемых штатных нормативов. Изменения по структуре {label} "
                  f"и штатной численности работников {label} могут вноситься по решению "
                  f"{leader_title_gen} {ORG_SHORT} на основании предложений начальника "
                  f"{label}, согласованных с начальником Управления.",
                  size=14, first_line_indent=indent)
    else:
        _add_para(doc,
                  f"1.{next_num}. Структура {label} и штатная численность работников {label} "
                  f"утверждаются {leader_title_ins} {ORG_SHORT} по предложению "
                  f"{head_appointment}, исходя из возложенных на {Label} функций и задач "
                  f"с учетом рекомендуемых штатных нормативов. Изменения по структуре {label} "
                  f"и штатной численности работников {label} могут вноситься по решению "
                  f"{leader_title_gen} {ORG_SHORT} на основании предложений "
                  f"начальника {label}.",
                  size=14, first_line_indent=indent)
    next_num += 1

    # Governing sources.
    if is_otdel:
        extra_docs = "положением об Управлении, а также настоящим Положением"
    else:
        extra_docs = "а также настоящим Положением"

    governing_text = ", ".join(governing_documents)
    _add_para(doc,
              f"1.{next_num}. {Label} в своей деятельности руководствуется "
              f"{governing_text}, {extra_docs}.",
              size=14, first_line_indent=indent)
    next_num += 1

    # Optional stamps.
    _add_para(doc,
              f"1.{next_num}. {Label} может иметь штампы, необходимые для осуществления "
              f"своей деятельности.",
              size=14, first_line_indent=indent)

    # Section 2: Goals and tasks.

    _add_para(doc, f"2. Основные цели и задачи {label}", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc, f"2.1. Основными целями {label} являются:",
              size=14, first_line_indent=indent)

    for i, goal in enumerate(goals, 1):
        _add_para(doc, f"2.1.{i}. {goal}", size=14, first_line_indent=indent)

    _add_para(doc, f"2.2. Основными задачами {label} являются:",
              size=14, first_line_indent=indent, space_before=6)

    for i, task in enumerate(tasks, 1):
        _add_para(doc, f"2.2.{i}. {task}", size=14, first_line_indent=indent)

    _add_para(doc,
              f"2.3. Цели и задачи {label} реализуются в пределах компетенции, "
              f"установленной настоящим Положением, и без дублирования функций иных "
              f"структурных подразделений {ORG_SHORT}.",
              size=14, first_line_indent=indent, space_before=6)

    # Section 3: Functions.

    _add_para(doc, f"3. Функции {label}", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc,
              f"В соответствии с поставленными задачами {Label} выполняет следующие функции:",
              size=14, first_line_indent=indent)

    for i, func in enumerate(functions, 1):
        _add_para(doc, f"3.{i}. {func}", size=14, first_line_indent=indent)

    # Section 4: Rights.

    _add_para(doc, f"4. Права {label}", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc, f"Для выполнения своих функций {Label} имеет право:",
              size=14, first_line_indent=indent)

    all_rights = _standard_rights(label, ORG_SHORT) + extra_rights
    for i, right in enumerate(all_rights, 1):
        _add_para(doc, f"4.{i}. {right}", size=14, first_line_indent=indent)

    # Section 5: Management and operations.

    _add_para(doc, f"5. Руководство и организация деятельности {label}", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc,
              f"5.1. Руководство {label} осуществляет начальник {label}, который "
              f"назначается на должность и освобождается от должности {leader_title_ins} "
              f"{ORG_SHORT} по представлению {head_appointment} "
              f"в соответствии с трудовым законодательством.",
              size=14, first_line_indent=indent)

    _add_para(doc, f"5.2. Начальник {label}:", size=14, first_line_indent=indent)

    if is_otdel:
        items_52 = [
            f"а) осуществляет непосредственное руководство деятельностью {label}, организует "
            f"его текущую работу по решению поставленных перед ним задач, несет персональную "
            f"ответственность за выполнение функций {label} и состояние исполнительской дисциплины;",

            f"б) докладывает начальнику Управления информацию по вопросам, входящим "
            f"в компетенцию {label}, а также вносит предложения по вопросам организации "
            f"и планирования деятельности {label};",

            f"в) принимает участие в разработке положения об {Label}е, его согласовании и утверждении;",

            f"г) разрабатывает проекты должностных инструкций работников {label};",

            "д) выполняет иные обязанности, возложенные на него должностной инструкцией.",
        ]
    else:
        items_52 = [
            f"а) осуществляет непосредственное руководство деятельностью {label}, организует "
            f"его текущую работу по решению поставленных перед ним задач, несет персональную "
            f"ответственность за выполнение функций {label} и состояние исполнительской дисциплины;",

            f"б) докладывает непосредственному руководителю информацию по вопросам, входящим "
            f"в компетенцию {label}, а также вносит предложения по вопросам организации "
            f"и планирования деятельности {label};",

            f"в) принимает участие в разработке положения об {Label}и, его согласовании и утверждении;",

            f"г) согласовывает проекты должностных инструкций работников отделов в {Label}и;",

            "д) выполняет иные обязанности, возложенные на него должностной инструкцией.",
        ]

    for item in items_52:
        _add_para(doc, item, size=14, first_line_indent=indent)

    _add_para(doc, f"5.3. Начальник {label} несет персональную ответственность за:",
              size=14, first_line_indent=indent, space_before=6)

    resp_items = [
        f"\u2013 организацию работы {label};",
        f"\u2013 своевременное и квалифицированное выполнение приказов, распоряжений, поручений "
        f"вышестоящего руководства, действующих нормативно-правовых актов по своему профилю деятельности;",
        "\u2013 рациональное и эффективное использование материальных, финансовых и кадровых ресурсов;",
        f"\u2013 состояние трудовой и исполнительской дисциплины в {Label}е, выполнение его работниками "
        f"своих функциональных обязанностей;" if is_otdel else
        f"\u2013 состояние трудовой и исполнительской дисциплины в {Label}и, выполнение его работниками "
        f"своих функциональных обязанностей;",
        f"\u2013 соблюдение работниками {label} правил внутреннего распорядка, "
        f"санитарно-противоэпидемического режима, противопожарной безопасности и техники безопасности;",
        "\u2013 ведение документации, предусмотренной действующим законодательством РФ;",
        f"\u2013 предоставление в установленном порядке достоверной статистической и иной информации "
        f"о деятельности {label}.",
    ]

    for item in resp_items:
        _add_para(doc, item, size=14, first_line_indent=indent)

    # 5.4
    _add_para(doc,
              f"5.4. Начальник {label} непосредственно подчиняется {head_reports_to}.",
              size=14, first_line_indent=indent, space_before=6)

    # 5.5
    _add_para(doc,
              f"5.5. В случае временного отсутствия начальника {label}, его обязанности "
              f"исполняются лицом, на которое по решению {leader_title_gen} {ORG_SHORT} "
              f"возлагается исполнение обязанностей начальника {label}.",
              size=14, first_line_indent=indent)

    # 5.6
    if is_otdel:
        _add_para(doc,
                  f"5.6. Работники {label} принимаются на работу по представлению начальника "
                  f"{label}, согласованному с начальником Управления, и осуществляют свою "
                  f"трудовую деятельность в соответствии с должностными инструкциями.",
                  size=14, first_line_indent=indent)
    else:
        _add_para(doc,
                  f"5.6. Работники {label} принимаются на работу по представлению начальников "
                  f"отделов, входящих в состав {label}, согласованному с начальником {label} "
                  f"и осуществляют свою трудовую деятельность в соответствии "
                  f"с должностными инструкциями.",
                  size=14, first_line_indent=indent)

    # 5.7
    _add_para(doc,
              f"5.7. Режим работы {label} определяется Правилами внутреннего трудового "
              f"распорядка Учреждения.",
              size=14, first_line_indent=indent)

    # Section 6: Responsibility.

    _add_para(doc, "6. Ответственность", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc,
              f"6.1. Начальник {label} несет ответственность за невыполнение функций, "
              f"предусмотренных настоящим Положением и правилами внутреннего распорядка "
              f"{ORG_SHORT}, а также за непринятие решений, входящих в сферу его компетенции.",
              size=14, first_line_indent=indent)

    if is_otdel:
        _add_para(doc,
                  f"6.2. Работники {label} несут ответственность за качество и своевременность "
                  f"выполнения должностных обязанностей, предусмотренных их должностными инструкциями.",
                  size=14, first_line_indent=indent)
    else:
        _add_para(doc,
                  f"6.2. Работники отделов, входящих в состав {label}, несут ответственность "
                  f"за качество и своевременность выполнения должностных обязанностей, "
                  f"предусмотренных их должностными инструкциями.",
                  size=14, first_line_indent=indent)

    _add_para(doc,
              f"6.3. Работники {label} несут ответственность за нарушение трудовой "
              f"дисциплины, правил техники безопасности при выполнении работ, за невыполнение "
              f"приказов и распоряжений руководства {ORG_SHORT} и начальника {label}, "
              f"входящих в сферу их компетенции.",
              size=14, first_line_indent=indent)

    # Section 7: Final provisions.

    _add_para(doc, "7. Заключительные положения", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)

    _add_para(doc,
              "7.1. Настоящее Положение вступает в силу с момента его утверждения "
              "и действует до замены новым.",
              size=14, first_line_indent=indent)

    # Sign-off table.

    _add_para(doc, "", size=14, space_before=12)
    _add_para(doc, "СОГЛАСОВАНО:", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    if approvers:
        table = doc.add_table(rows=len(approvers) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_borders(table)

        headers = ["Должность", "Ф.И.О.", "Подпись", "Дата"]
        for i, h in enumerate(headers):
            _cell_para(table.rows[0].cells[i], h, bold=True, size=12)

        for idx, approver in enumerate(approvers, 1):
            _cell_para(table.rows[idx].cells[0], approver["position"], size=12)
            _cell_para(table.rows[idx].cells[1], approver.get("name", ""), size=12)

    # Acquaintance sheet.

    doc.add_page_break()
    _add_para(doc, "Лист ознакомления", bold=True, size=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    if is_otdel:
        title_gen = (
            f"с положением об отделе {unit_name} {parent_unit_name} {ORG_FULL}"
        )
    else:
        title_gen = (
            f"с положением об Управлении {unit_name} {ORG_FULL}"
        )

    _add_para(doc, title_gen, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    _add_para(doc, "Ознакомлен(а):", size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

    oz_table = doc.add_table(rows=15, cols=3)
    oz_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(oz_table)
    headers_oz = ["Ф.И.О.", "Дата", "Подпись"]
    for i, h in enumerate(headers_oz):
        _cell_para(oz_table.rows[0].cells[i], h, bold=True, size=12)

    output_path = _resolve_output_path(
        output_path or _default_output_path(org, unit_type, unit_name)
    )
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    print(f"Сохранено: {output_path}")
    return output_path
