#!/usr/bin/env python3
"""Генератор протоколов совещаний (.docx).

Конфигурация организации загружается из ~/.docs-plugin/org_details.md.
При отсутствии файла запустите docs-init для первичной настройки.

Использование:
    python3 generate.py  # создаёт пример протокола
"""

import os
import shutil
import warnings as _warnings
import re as _re
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys as _sys

_HERE = os.path.dirname(os.path.abspath(__file__))
for _lib_path in (
    os.path.join(_HERE, "..", "..", "lib"),
    os.path.expanduser("~/.docs-plugin/runtime/lib"),
):
    if os.path.isdir(_lib_path) and _lib_path not in _sys.path:
        _sys.path.insert(0, _lib_path)
from db import log_generation, ORG_DETAILS_PATH
from docx_meta import new_document


_SLASH_OK = _re.compile(r'п/п|ИНН/КПП|[0-9]/[0-9]|[A-Za-z]/[A-Za-z]')


def _warn_slash(text):
    """Предупреждение при обнаружении косой черты между кириллическими словами."""
    for m in _re.finditer(r'[а-яА-ЯёЁ]+\s*/\s*[а-яА-ЯёЁ]+', text):
        if not _SLASH_OK.search(m.group()):
            _warnings.warn(
                f"Косая черта в тексте документа: «{m.group()}». "
                f"Используйте запятую, «или», «и» или скобки.",
                stacklevel=3,
            )


def _load_org_details():
    """Загрузить реквизиты из ~/.docs-plugin/org_details.md."""
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        print("WARNING: ~/.docs-plugin/org_details.md not found. Run docs-init to configure.")
        return config
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()
    return config


def _load_staff():
    """Загрузить штатный список из staff_file (xlsx).

    Возвращает list[{"lastname", "initials", "position"}].
    При отсутствии staff_file возвращает [] (сверка станет no-op).
    """
    cfg = _load_org_details()
    path = os.path.expanduser(cfg.get("staff_file", "") or "")
    if not path or not os.path.exists(path):
        return []
    try:
        from openpyxl import load_workbook
    except ImportError:
        return []
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    header = [str(c or "").strip().lower() for c in rows[0]]
    def idx(name_options):
        for opt in name_options:
            if opt in header:
                return header.index(opt)
        return None
    i_fio = idx(["фио", "ф.и.о.", "сотрудник"])
    i_pos = idx(["должность"])
    if i_fio is None:
        return []
    out = []
    for row in rows[1:]:
        cell = row[i_fio] if i_fio < len(row) else None
        if not cell:
            continue
        s = str(cell).strip()
        # «Фамилия Имя Отчество» или «Фамилия И.О.»
        parts = s.split()
        if len(parts) >= 2:
            lastname = parts[0]
            if len(parts) == 2 and "." in parts[1]:
                initials = parts[1]
            else:
                initials = ".".join(p[0] for p in parts[1:3]) + "."
            position = ""
            if i_pos is not None and i_pos < len(row) and row[i_pos]:
                position = str(row[i_pos]).strip().lower()
            out.append({"lastname": lastname, "initials": initials, "position": position})
    return out


# ─── Вспомогательные функции ──────────────────────────────────────────────────

def _set_cell_margins_zero(table):
    """Установить внутренние отступы ячеек таблицы в 0 со всех сторон."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def _remove_table_borders(table):
    """Убрать все границы таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _autofit_contents(table):
    """Установить AutoFit Contents для таблицы.

    Таблица занимает 100% ширины листа, столбцы подстраиваются под содержимое.
    Применять ко всем содержательным таблицам (перечни, листы ознакомления и т.д.).
    Не применять к layout-таблицам (подпись, дата/номер).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Ширина таблицы 100% (5000 = 100% в пятидесятых долях процента)
    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    # tblLayout autofit
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)
    # Ширины столбцов → auto
    for row in tbl.findall(qn("w:tr")):
        for tc in row.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is not None:
                    tcW.set(qn("w:type"), "auto")


def _set_rows_no_break(table):
    """Запретить разрыв строк таблицы между страницами (cantSplit)."""
    for row in table.rows:
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "true")
        trPr.append(cantSplit)


def _set_table_borders(table):
    """Установить тонкие границы для таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _run(para, text, size=14, bold=False, italic=False, font_name="Times New Roman"):
    """Добавить run к параграфу с настройками шрифта и языком ru-RU.

    Если *text* содержит ``\\n``, каждый перенос превращается в мягкий
    разрыв строки (``<w:br/>``, аналог Shift+Enter в Word).
    """
    text = text.replace("—", "–")  # длинное тире -> среднее
    _warn_slash(text)
    parts = text.split("\n")
    run = None
    for i, part in enumerate(parts):
        if i > 0:
            br = OxmlElement("w:br")
            run._element.append(br)
        if part or i == 0:
            run = para.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            rPr = run._element.rPr
            rPr.rFonts.set(qn("w:eastAsia"), font_name)
            lang = rPr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rPr.append(lang)
            lang.set(qn("w:val"), "ru-RU")
            lang.set(qn("w:eastAsia"), "ru-RU")
            lang.set(qn("w:bidi"), "ru-RU")
    return run


def _para(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, space_after=None, space_before=None):
    """Добавить параграф с настройками."""
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    if text:
        _run(p, text, size=size, bold=bold)
    return p


def _dash_para(doc, text, size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Буллет-пункт с «–» (средним тире). Отступ слева 1.25 см, без красной строки."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    _run(p, f"– {text}", size=size)
    return p


def _cell_text(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               vertical_align="center"):
    """Установить текст ячейки с форматированием.

    vertical_align: "center" | "top" | "bottom" | None.
    По умолчанию "center" – вертикальное выравнивание посередине.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if text:
        _run(para, text, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _cell_lines(cell, lines, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_align="center"):
    """Установить несколько строк в ячейке как отдельные параграфы (без \\n).

    lines: list[str] – каждая строка станет отдельным параграфом.
    """
    cell.text = ""
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        if line:
            _run(para, line, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _build_org_config():
    """Конфиг организации из org_details.md + output_dir_protocol."""
    o = _load_org_details()

    # 4 строки шапки: parent_org (или parent_org_short если задан),
    # full_name, "(short_name)", пустая строка-разделитель.
    header_lines = []
    parent = o.get("parent_org_short") or o.get("parent_org")
    if parent:
        header_lines.append(parent)
    if o.get("full_name"):
        header_lines.append(o["full_name"])
    if o.get("short_name"):
        header_lines.append(f"({o['short_name']})")

    return {
        "header_lines": header_lines,
        "address": o.get("address", ""),
        "phone_fax": o.get("phone_fax", ""),
        "leader_title": o.get("leader_title", ""),
        "leader_name_nom": o.get("leader_name_nom", ""),
        "short_name": o.get("short_name", ""),
        "full_name": o.get("full_name", ""),
        "staff_file": o.get("staff_file", ""),
        "output_dir_protocol": o.get("output_dir_protocol", ""),
    }


_RU_MONTHS = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def _format_ru_date(date_str):
    """«DD.MM.YYYY» → «D месяца YYYY г.» (день без ведущего нуля)."""
    d, m, y = date_str.split(".")
    return f"{int(d)} {_RU_MONTHS[int(m) - 1]} {y} г."


def _format_fio_initials_first(fio):
    """«Фамилия И.О.» → «И.О. Фамилия». Если уже в нужном порядке — вернуть как есть."""
    fio = fio.strip()
    parts = fio.split()
    if len(parts) != 2:
        return fio
    a, b = parts
    # Эвристика: инициалы — две заглавные с точками.
    if "." in a and "." not in b:
        return fio  # уже «И.О. Фамилия»
    return f"{b} {a}"


def _make_header_block(doc, org):
    """Шапка: 4 строки центр+жирный 14pt + пустая + 2 строки 10pt."""
    for line in org["header_lines"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, line, bold=True, size=14)
    # Пустая строка-разделитель
    doc.add_paragraph()
    if org["address"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, org["address"], italic=True, size=10)
    if org["phone_fax"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, org["phone_fax"], italic=True, size=10)


def _make_title_block(doc, subtype):
    """«ПРОТОКОЛ\n<подтип>» — две строки, центр+жирный."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, "ПРОТОКОЛ", bold=True, size=14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, subtype, bold=True, size=14)


def _make_date_number_table(doc, doc_date, doc_number):
    """Таблица 1×2: «12 мая 2026 г.» (слева) | «№ XXX» или «№ _________» (справа)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)
    _set_cell_margins_zero(table)
    left_text = _format_ru_date(doc_date)
    right_text = f"№ {doc_number}" if doc_number else "№ _________"
    _cell_text(table.rows[0].cells[0], left_text)
    _cell_text(table.rows[0].cells[1], right_text)
    # Правая ячейка — выравнивание вправо
    for p in table.rows[0].cells[1].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table


def _make_venue_chair_block(doc, venue, chair):
    """Три параграфа: «Место проведения: ...», «Председатель: ...» (метка жирная)."""
    if venue:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        _run(p, "Место проведения: ", bold=True, size=14)
        _run(p, venue, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    _run(p, "Председатель: ", bold=True, size=14)
    chair_text = f"{chair['lastname']} {chair['initials']} — {chair['position']}"
    _run(p, chair_text, size=14)


def _make_attendees_table(doc, attendees, chair):
    """Параграф-метка «Присутствовали:» + 3-col таблица: ФИО | – | должность.

    Председатель удаляется из attendees с предупреждением — он не дублируется в списке.
    """
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label.paragraph_format.first_line_indent = Cm(1.25)
    _run(label, "Присутствовали:", bold=True, size=14)

    key = (chair["lastname"].strip().lower(), chair["initials"].strip().lower())
    filtered = []
    for a in attendees:
        if (a["lastname"].strip().lower(), a["initials"].strip().lower()) == key:
            _warnings.warn(
                f"Председатель «{chair['lastname']} {chair['initials']}» "
                f"передан в attendees — отфильтрован.",
                stacklevel=2,
            )
            continue
        filtered.append(a)

    table = doc.add_table(rows=len(filtered), cols=3)
    _remove_table_borders(table)
    _set_cell_margins_zero(table)
    _autofit_contents(table)
    for i, a in enumerate(filtered):
        cells = table.rows[i].cells
        _cell_text(cells[0], f"{a['lastname']} {a['initials']}")
        _cell_text(cells[1], "–")
        _cell_text(cells[2], a["position"])
    return table


def _next_num_id(doc):
    """Вернуть следующий свободный numId в numbering.xml."""
    try:
        numbering = doc.part.numbering_part.element
    except (AttributeError, ValueError):
        # Если numbering_part не создан — создать через временный список
        tmp = doc.add_paragraph(style="List Number")
        tmp._p.getparent().remove(tmp._p)
        numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    existing = [int(n.get(qn("w:numId"))) for n in nums]
    return (max(existing) + 1) if existing else 1


def _setup_protocol_numbering(doc):
    """Зарегистрировать decimal-нумерацию «%1.» с отступами left=1069 EMU, hanging=360.

    Возвращает numId (динамический, не хардкод).
    """
    num_id = _next_num_id(doc)
    numbering = doc.part.numbering_part.element

    # Найти следующий свободный abstractNumId
    abs_nums = numbering.findall(qn("w:abstractNum"))
    abs_id = (max(int(n.get(qn("w:abstractNumId"))) for n in abs_nums) + 1) if abs_nums else 0

    from docx.oxml import parse_xml
    abs_xml = (
        f'<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{abs_id}">'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="decimal"/>'
        '<w:lvlText w:val="%1."/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="1069" w:hanging="360"/></w:pPr>'
        '</w:lvl>'
        '</w:abstractNum>'
    )
    numbering.append(parse_xml(abs_xml))

    num_xml = (
        f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        '</w:num>'
    )
    numbering.append(parse_xml(num_xml))
    return num_id


def _setup_dash_numbering(doc):
    """En-dash bullet для подпунктов. Возвращает numId."""
    num_id = _next_num_id(doc)
    numbering = doc.part.numbering_part.element
    abs_nums = numbering.findall(qn("w:abstractNum"))
    abs_id = (max(int(n.get(qn("w:abstractNumId"))) for n in abs_nums) + 1) if abs_nums else 0
    from docx.oxml import parse_xml
    abs_xml = (
        f'<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{abs_id}">'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="–"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="1429" w:hanging="360"/></w:pPr>'
        '</w:lvl>'
        '</w:abstractNum>'
    )
    numbering.append(parse_xml(abs_xml))
    num_xml = (
        f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        '</w:num>'
    )
    numbering.append(parse_xml(num_xml))
    return num_id


def _attach_numbering(para, num_id, ilvl=0):
    """Присоединить numPr к параграфу."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_el = OxmlElement("w:numId")
    num_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_el)
    pPr.append(numPr)


def create_protocol(*args, **kwargs):
    """Заглушка — реализация в Task B15."""
    raise NotImplementedError("create_protocol будет реализован в Task B15")
