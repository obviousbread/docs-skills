#!/usr/bin/env python3
"""Генератор официальных писем.

Создает .docx файл с угловым бланком, телом письма, подписью и футером исполнителя.

Конфигурация организации загружается из references/org_details.md.
При отсутствии файла запустите docs-init для первичной настройки.

Использование:
    python3 generate.py  # создаёт пример письма
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import re as _re
import warnings as _warnings

_SLASH_OK = _re.compile(r'п/п|ИНН/КПП|[0-9]/[0-9]|[A-Za-z]/[A-Za-z]')


def _warn_slash(text):
    """Предупреждение при обнаружении косой черты между кириллическими словами."""
    for m in _re.finditer(r'[а-яА-ЯёЁ]+\s*/\s*[а-яА-ЯёЁ]+', text):
        if not _SLASH_OK.search(m.group()):
            _warnings.warn(
                f"Косая черта в тексте документа: \u00ab{m.group()}\u00bb. "
                f"Используйте запятую, \u00abили\u00bb, \u00abи\u00bb или скобки.",
                stacklevel=3,
            )

import sys as _sys

_HERE = os.path.dirname(os.path.abspath(__file__))
for _lib_path in (
    os.path.join(_HERE, "..", "..", "lib"),
    os.path.expanduser("~/.docs-plugin/runtime/lib"),
):
    if os.path.isdir(_lib_path) and _lib_path not in _sys.path:
        _sys.path.insert(0, _lib_path)
from db import log_generation, ORG_DETAILS_PATH


# ─── Загрузка реквизитов из org_details.md ────────────────────────────────────

def _load_org_details():
    """Загрузить реквизиты из ~/.docs-plugin/org_details.md."""
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        print("WARNING: ~/.docs-plugin/org_details.md not found. Run docs-init to configure.")
        return config
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()
    return config


def _build_org_config():
    """Построить конфиг организации из org_details.md."""
    o = _load_org_details()

    full_name = o.get("full_name", "")
    short_name = o.get("short_name", "")

    full_name_lines = []
    if full_name:
        full_name_lines.append(full_name)
    if short_name:
        full_name_lines.append(f"({short_name})")

    address_lines = []
    if o.get("address"):
        address_lines.append(f'адрес:{o["address"]}')
    if o.get("phone_fax"):
        address_lines.append(o["phone_fax"])
    if o.get("email"):
        address_lines.append(f'e-mail: {o["email"]}')
    if o.get("okpo") and o.get("ogrn"):
        address_lines.append(f'ОКПО, ОГРН {o["okpo"]}/{o["ogrn"]}')
    if o.get("inn_kpp"):
        address_lines.append(f'ИНН/КПП {o["inn_kpp"]}')

    return {
        "parent_org": o.get("parent_org", ""),
        "parent_org_short": o.get("parent_org_short", ""),
        "full_name_lines": full_name_lines,
        "address_lines": address_lines,
        "org_short": short_name,
        "org_full_text": full_name,
        "leader_title": o.get("leader_title", ""),
        "leader_name": o.get("leader_name_nom", ""),
        "executor_name": o.get("author_name_full", ""),
        "executor_phone": o.get("author_phone", ""),
    }


# ─── Вспомогательные функции ─────────────────────────────────────────────────

def _remove_table_borders(table):
    """Убрать все границы таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Удаляем существующие borders если есть
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    # Вставляем перед tblLook (OOXML требует строгий порядок элементов)
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblLook is not None:
        tblLook.addprevious(tblBorders)
    else:
        tblPr.append(tblBorders)


def _set_table_indent(table, indent_twips):
    """Установить отступ таблицы (может быть отрицательным)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(existing)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(indent_twips))
    tblInd.set(qn("w:type"), "dxa")
    # Вставляем перед tblBorders или tblLook (OOXML порядок)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblBorders is not None:
        tblBorders.addprevious(tblInd)
    elif tblLook is not None:
        tblLook.addprevious(tblInd)
    else:
        tblPr.append(tblInd)


def _set_table_width(table, width_twips):
    """Установить ширину таблицы в DXA."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    # tblW должен идти перед jc, tblInd, tblBorders, tblLook
    # Вставляем после tblStyle (если есть) или в начало
    tblStyle = tblPr.find(qn("w:tblStyle"))
    if tblStyle is not None:
        tblStyle.addnext(tblW)
    else:
        tblPr.insert(0, tblW)


def _set_row_height(row, height_twips):
    """Установить высоту строки таблицы."""
    trPr = row._tr.get_or_add_trPr()
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trPr.append(trHeight)


def _set_cell_width(cell, width_twips):
    """Установить ширину ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_grid_cols(table, widths):
    """Установить gridCol для таблицы."""
    tbl = table._tbl
    for existing in tbl.findall(qn("w:tblGrid")):
        tbl.remove(existing)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    # Вставляем после tblPr
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _run(para, text, size=14, bold=False, underline=False,
         font_name="Times New Roman"):
    """Добавить run к параграфу с настройками шрифта."""
    _warn_slash(text)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return run


def _para(doc, text, size=14, bold=False, underline=False,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, left_indent=None,
          space_after=None, space_before=None,
          line_spacing=1.0):
    """Добавить параграф с настройками."""
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        _run(p, text, size=size, bold=bold, underline=underline)
    return p


def _cell_para(cell, text, size=14, bold=False, underline=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, left_indent_twips=0):
    """Добавить параграф в ячейку."""
    if not cell.paragraphs or cell.paragraphs[0].text:
        para = cell.add_paragraph()
    else:
        para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if left_indent_twips:
        para.paragraph_format.left_indent = Twips(left_indent_twips)
    if text:
        _run(para, text, size=size, bold=bold, underline=underline)
    return para


def _cell_add_para(cell, text, size=14, bold=False, underline=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT, left_indent_twips=0):
    """Добавить дополнительный параграф в ячейку (всегда новый)."""
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if left_indent_twips:
        para.paragraph_format.left_indent = Twips(left_indent_twips)
    if text:
        _run(para, text, size=size, bold=bold, underline=underline)
    return para


def _set_contextual_spacing(para):
    """Включить contextualSpacing для параграфа."""
    pPr = para._element.get_or_add_pPr()
    cs = OxmlElement("w:contextualSpacing")
    pPr.append(cs)


# ─── Построение шапки-бланка ─────────────────────────────────────────────────

def _build_header_table(doc, date_str, number_str, on_number_str,
                        addressee_lines):
    """Создать угловую таблицу-шапку бланка.

    Параметры
    ---------
    doc : Document
    date_str : str
        Дата документа, например "30.07.2025"
    number_str : str
        Номер документа, например "01-1031"
    on_number_str : str
        «На №» (для ответных писем), например "32-024/13-3/1217 от 18.07.2025".
        Пустая строка если инициативное письмо.
    addressee_lines : list[str]
        Строки адресата, например:
        ["Генеральному директору", "ФГБУ «Научный центр»",
         "", "[И.О. Фамилия]"]
    """
    cfg = _build_org_config()

    # Таблица: 3 колонки (левая реквизиты | разделитель | правая адресат)
    table = doc.add_table(rows=1, cols=3)
    _remove_table_borders(table)
    _set_table_width(table, 10310)
    _set_table_indent(table, -671)
    _set_grid_cols(table, [4786, 709, 4815])
    _set_row_height(table.rows[0], 4819)

    # Ширина ячеек
    _set_cell_width(table.cell(0, 0), 4786)
    _set_cell_width(table.cell(0, 1), 709)
    _set_cell_width(table.cell(0, 2), 4815)

    left_cell = table.cell(0, 0)
    spacer_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    # ── Левая колонка: реквизиты ──────────────────────────────────────────

    # Надзорный орган (parent_org) — 11pt bold center, если указан
    _first_written = False
    if cfg.get("parent_org"):
        display_text = cfg.get("parent_org_short") or cfg["parent_org"].upper()
        _cell_para(left_cell, display_text, size=11, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_add_para(left_cell, "", size=11,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _first_written = True

    # Полное наименование — 11pt bold center
    for line in cfg["full_name_lines"]:
        if not _first_written:
            _cell_para(left_cell, line, size=11, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _first_written = True
        else:
            _cell_add_para(left_cell, line, size=11, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER)

    # Пустая строка
    _cell_add_para(left_cell, "", size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Адрес и реквизиты — 11pt regular center
    for line in cfg["address_lines"]:
        _cell_add_para(left_cell, line, size=11,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    # Пустая строка
    _cell_add_para(left_cell, "", size=14,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Дата и номер — 14pt underlined center
    if date_str and number_str:
        date_num_text = f"от {date_str} № {number_str}"
    elif date_str:
        date_num_text = f"от {date_str} №"
    elif number_str:
        date_num_text = f"от ____________ № {number_str}"
    else:
        date_num_text = "от ____________ №___________"
    _cell_add_para(left_cell, date_num_text, size=14, underline=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # «на ___» — 14pt underlined center
    if on_number_str:
        on_text = f"на {on_number_str}"
    else:
        on_text = "на _________________"
    _cell_add_para(left_cell, on_text, size=14, underline=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Финальный пустой параграф в левой ячейке
    _cell_add_para(left_cell, "", size=14,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Средняя колонка: пусто ────────────────────────────────────────────
    spacer_cell.paragraphs[0].text = ""

    # ── Правая колонка: адресат ───────────────────────────────────────────
    first = True
    for line in addressee_lines:
        if first:
            _cell_para(right_cell, line, size=14,
                       align=WD_ALIGN_PARAGRAPH.LEFT,
                       left_indent_twips=421)
            first = False
        else:
            _cell_add_para(right_cell, line, size=14,
                           align=WD_ALIGN_PARAGRAPH.LEFT,
                           left_indent_twips=421)

    return table


# ─── Построение блока подписи ────────────────────────────────────────────────

def _build_signature_table(doc, signer_title, signer_name):
    """Создать таблицу подписи (2 колонки без границ, 100% ширины)."""
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)

    # Установить ширину таблицы 100%
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    # Вставляем в начало (перед tblLook)
    tblPr.insert(0, tblW)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Должность — по левому краю
    _cell_para(left_cell, signer_title, size=14,
               align=WD_ALIGN_PARAGRAPH.LEFT)

    # ФИО — по правому краю
    _cell_para(right_cell, signer_name, size=14,
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    return table


# ─── Построение футера (исполнитель) ─────────────────────────────────────────

def _build_first_page_footer(section, executor_name, executor_phone):
    """Настроить футер первой страницы с данными исполнителя."""
    # Включить «Особый колонтитул первой страницы»
    section.different_first_page_header_footer = True

    footer = section.first_page_footer
    footer.is_linked_to_previous = False

    # Имя исполнителя
    p1 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.line_spacing = 1.0
    run1 = p1.add_run(executor_name)
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(8)
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Телефон
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(executor_phone)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(8)
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _build_executor_body(doc, executor_name, executor_title, executor_phone):
    """Добавить исполнителя как обычный текст в теле документа (серый, 8pt).

    Используется для многостраничных писем, где колонтитул первой страницы
    не подходит - исполнитель должен быть на той же странице, что и подписант.

    Формат:
        Фамилия Имя Отчество,
        должность, ученая степень,
        телефон
    """
    GRAY = RGBColor(0x80, 0x80, 0x80)

    lines = []
    if executor_title:
        lines.append(f"{executor_name},")
        lines.append(f"{executor_title},")
        lines.append(executor_phone)
    else:
        lines.append(executor_name)
        lines.append(executor_phone)

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


# ─── Автопунктуация маркированных списков ─────────────────────────────────────

def _auto_punctuate_dashed(paragraphs):
    """Расставить «;» после каждого dashed-элемента и «.» после последнего в группе.

    Работает с копией списка, не мутирует оригинал. Не трогает элементы,
    у которых текст уже заканчивается на знак препинания (.;:!?).
    """
    items = []
    for p in paragraphs:
        if isinstance(p, dict):
            items.append(dict(p))  # shallow copy
        else:
            items.append(p)

    # Найти группы подряд идущих dashed-элементов
    i = 0
    while i < len(items):
        if isinstance(items[i], dict) and items[i].get("dashed"):
            # Начало группы
            group_start = i
            while i < len(items) and isinstance(items[i], dict) and items[i].get("dashed"):
                i += 1
            group_end = i  # exclusive
            # Расставить пунктуацию
            for j in range(group_start, group_end):
                text = items[j].get("text", "")
                if text and text[-1] not in ".;:!?":
                    if j == group_end - 1:
                        items[j]["text"] = text + "."
                    else:
                        items[j]["text"] = text + ";"
        else:
            i += 1

    return items


# ─── Основная функция ────────────────────────────────────────────────────────

def create_letter(
    addressee_lines=None,
    greeting="Уважаемые коллеги!",
    body_paragraphs=None,
    appendix_text="",
    signer_title=None,
    signer_name=None,
    doc_date="",
    doc_number="",
    on_number="",
    executor_name=None,
    executor_title="",
    executor_phone=None,
    executor_in_body=False,
    output_path=None,
):
    """
    Создает официальное письмо на угловом бланке организации в формате .docx.

    Параметры
    ---------
    addressee_lines : list[str]
        Строки адресата. Каждая строка — отдельный абзац в правой колонке.
        Пример: ["Генеральному директору",
                 "ФГБУ «Научный центр»", "", "[И.О. Фамилия]"]
        Пустые строки создают визуальный разрыв.

    greeting : str
        Обращение. По умолчанию "Уважаемые коллеги!".
        Примеры: "Уважаемый [Имя Отчество]!", "Уважаемая [Имя Отчество]!"

    body_paragraphs : list[str | dict]
        Абзацы тела письма. Каждый элемент — один абзац с красной строкой.
        Строка → обычный абзац.
        dict с ключами:
            "text" : str — текст
            "indent" : bool — красная строка (по умолчанию True)
            "numbered" : bool — нумерованный пункт (1., 2., ...)
            "dashed" : bool — буллет с «–»
            "bold" : bool — жирный текст (для подзаголовков)
            "center" : bool — по центру

    appendix_text : str
        Текст приложения. Пример: "Извещение на 1 л. в 1 экз."
        Формируется как "Приложение: <text>".
        Пустая строка — приложение не добавляется.

    signer_title : str
        Должность подписанта. По умолчанию из конфигурации.

    signer_name : str
        ФИО подписанта (И.О. Фамилия). По умолчанию из конфигурации.

    doc_date : str
        Дата документа "ДД.ММ.ГГГГ". Пустая строка — подчёркивание.

    doc_number : str
        Номер документа. Пустая строка — подчёркивание.

    on_number : str
        «На №» — для ответных писем. Например: "№ 32-024/13-3/1217 от 18.07.2025".
        Пустая строка — стандартное подчёркивание.

    executor_name : str
        ФИО исполнителя (Фамилия Имя Отчество). По умолчанию из конфигурации.

    executor_title : str
        Должность исполнителя. Используется при executor_in_body=True.
        Пример: "начальник Управления по перспективному развитию медицинской деятельности, к.м.н."

    executor_phone : str
        Телефон исполнителя. По умолчанию из конфигурации.

    executor_in_body : bool
        Если True — исполнитель размещается в теле документа серым текстом (8pt)
        на странице с подписантом, а не в колонтитуле первой страницы.
        Использовать для многостраничных писем (>1 страницы).
        По умолчанию False (колонтитул первой страницы).

    output_path : str | None
        Путь для сохранения. По умолчанию "Письмо_<дата>.docx".

    Returns
    -------
    str : путь к сохранённому файлу
    """
    cfg = _build_org_config()

    if addressee_lines is None:
        addressee_lines = ["[Адресат]"]
    if body_paragraphs is None:
        body_paragraphs = ["[Текст письма]"]
    if signer_title is None:
        signer_title = cfg["leader_title"]
    if signer_name is None:
        signer_name = cfg["leader_name"]
    if executor_name is None:
        executor_name = cfg["executor_name"]
    if executor_phone is None:
        executor_phone = cfg["executor_phone"]

    if output_path is None:
        d = doc_date if doc_date else date.today().strftime("%d.%m.%Y")
        output_path = f"Письмо {d}.docx"

    doc = Document()

    # ── Настройка страницы (A4) ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Emu(11900 * 635)   # 11900 DXA → EMU
    section.page_height = Emu(16840 * 635)  # 16840 DXA → EMU
    section.top_margin = Emu(1134 * 635)    # ~2 см
    section.bottom_margin = Emu(1134 * 635) # ~2 см
    section.left_margin = Emu(1701 * 635)   # ~3 см
    section.right_margin = Emu(851 * 635)   # ~1.5 см
    section.header_distance = Emu(720 * 635)
    section.footer_distance = Emu(720 * 635)

    # Шрифт по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    # ── Шапка-бланк ──────────────────────────────────────────────────────
    _build_header_table(doc, doc_date, doc_number, on_number,
                        addressee_lines)

    # ── Пустые строки после шапки ────────────────────────────────────────
    _para(doc, "", space_after=0)
    _para(doc, "", space_after=0)
    _para(doc, "", space_after=0)

    # ── Обращение ────────────────────────────────────────────────────────
    _para(doc, greeting, size=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    # Пустая строка после обращения
    _para(doc, "", space_after=0)

    # ── Тело письма ──────────────────────────────────────────────────────
    num_counter = [0]  # mutable для нумерации

    # Автопунктуация dashed-элементов: «;» после каждого, «.» после последнего в группе
    body_paragraphs = _auto_punctuate_dashed(body_paragraphs)

    for item in body_paragraphs:
        if isinstance(item, str):
            # Обычный абзац с красной строкой
            p = _para(doc, item, size=14,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      first_indent=1.25, space_after=0)
        elif isinstance(item, dict):
            text = item.get("text", "")
            is_bold = item.get("bold", False)
            is_center = item.get("center", False)
            is_numbered = item.get("numbered", False)
            is_dashed = item.get("dashed", False)
            has_indent = item.get("indent", True)

            if is_numbered:
                num_counter[0] += 1
                text = f"{num_counter[0]}. {text}"
                p = _para(doc, text, size=14, bold=is_bold,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_indent=1.25, space_after=0)
            elif is_dashed:
                text = f"– {text}"
                p = _para(doc, text, size=14, bold=is_bold,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_indent=1.25, space_after=0)
            elif is_center:
                p = _para(doc, text, size=14, bold=is_bold,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=0)
            else:
                fi = 1.25 if has_indent else None
                p = _para(doc, text, size=14, bold=is_bold,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_indent=fi, space_after=0)

    # ── Пустая строка после текста ───────────────────────────────────────
    _para(doc, "", space_after=0)

    # ── Приложение ───────────────────────────────────────────────────────
    if appendix_text:
        p = _para(doc, "", size=14,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=1.25, space_after=0)
        _run(p, "Приложение: ", size=14)
        _run(p, appendix_text, size=14)

    # ── Пустые строки перед подписью ─────────────────────────────────────
    _para(doc, "", space_after=0)
    _para(doc, "", space_after=0)

    # ── Подпись ──────────────────────────────────────────────────────────
    _build_signature_table(doc, signer_title, signer_name)

    # ── Исполнитель ──────────────────────────────────────────────────────
    if executor_in_body:
        # Многостраничное письмо: исполнитель в теле документа (серый, 8pt)
        for _ in range(7):
            _para(doc, "", space_after=0)
        _build_executor_body(doc, executor_name, executor_title, executor_phone)
    else:
        # Одностраничное письмо: исполнитель в колонтитуле первой страницы
        for _ in range(7):
            _para(doc, "", space_after=0)
        _build_first_page_footer(section, executor_name, executor_phone)

    # ── Сохранение ───────────────────────────────────────────────────────
    output_path = os.path.expanduser(output_path)
    doc.save(output_path)
    log_generation("letter", greeting, output_path, {
        "has_appendix": bool(appendix_text),
        "executor_in_body": executor_in_body,
        "has_on_number": bool(on_number),
    })
    return output_path


# ─── Демонстрация ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    path = create_letter(
        addressee_lines=[
            "Отдел бухгалтерского учета",
            "и отчетности",
            "ФГБУ «Научный центр»",
        ],
        greeting="Уважаемые коллеги!",
        body_paragraphs=[
            (
                "Федеральное государственное бюджетное учреждение "
                "«Центр тестирования» направляет подписанное "
                "и заверенное печатью извещение о передаче лекарственных "
                "препаратов по централизованным поставкам (2025 год)."
            ),
        ],
        appendix_text="Извещение на 1 л. в 1 экз.",
        doc_date="30.07.2025",
        doc_number="01-1031",
        on_number="",
        output_path="Пример письма.docx",
    )
    print(f"Создано: {path}")
