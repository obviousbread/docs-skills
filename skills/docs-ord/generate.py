#!/usr/bin/env python3
"""Генератор организационно-распорядительных документов (ОРД).

Типы документов: приказ, распоряжение, указание.

Конфигурация организации загружается из references/org_details.md.
При отсутствии файла запустите docs-init для первичной настройки.

Использование:
    python3 generate.py  # создаёт пример приказа
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

import re as _re
import warnings as _warnings

_SLASH_OK = _re.compile(r'п/п|ИНН/КПП|[0-9]/[0-9]|[A-Za-z]/[A-Za-z]')

def _warn_slash(text):
    """Предупреждение при обнаружении косой черты между кириллическими словами."""
    for m in _re.finditer(r'[а-яА-ЯёЁ]+\s*/\s*[а-яА-ЯёЁ]+', text):
        if not _SLASH_OK.search(m.group()):
            _warnings.warn(
                f"Косая черта в тексте документа: \u00ab{m.group()}\u00bb. "
                f"Используйте запятую, \u00abили\u00bb, \u00abи\u00bb или скобки.",
                stacklevel=3,
            )

import sys as _sys

_HERE = os.path.dirname(os.path.abspath(__file__))
for _lib_path in (
    os.path.join(_HERE, "..", "..", "lib"),
    os.path.expanduser("~/.docs-plugin/runtime/lib"),
):
    if os.path.isdir(_lib_path) and _lib_path not in _sys.path:
        _sys.path.insert(0, _lib_path)
from db import log_generation, ORG_DETAILS_PATH


# ─── Загрузка реквизитов из org_details.md ────────────────────────────────────

def _load_org_details():
    """Загрузить реквизиты из ~/.docs-plugin/org_details.md."""
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        print("WARNING: ~/.docs-plugin/org_details.md not found. Run docs-init to configure.")
        return config
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()
    return config


def _build_org_config():
    """Построить конфиг организации из org_details.md."""
    o = _load_org_details()

    header_lines = []
    if o.get("parent_org"):
        header_lines.append((o["parent_org"], True, False, True, 14))
    if o.get("full_name"):
        header_lines.append((o["full_name"], True, False, False, 14))
    if o.get("short_name"):
        header_lines.append(("(" + o["short_name"] + ")", True, False, False, 14))
    header_lines.append(("", False, False, False))
    if o.get("address"):
        header_lines.append((o["address"], True, True, False, 10))
    if o.get("phone_fax"):
        header_lines.append((o["phone_fax"], True, True, True, 10))

    approvers = []
    for i in range(1, 4):
        name = o.get(f"approver_{i}_name", "")
        position = o.get(f"approver_{i}_position", "")
        if name and position:
            approvers.append({"name": name, "position": position})

    return {
        "header_font_size": 11,
        "header_lines": header_lines,
        "city": None,
        "leader_title": o.get("leader_title", ""),
        "leader_name": o.get("leader_name_nom", ""),
        "keyword_upper": True,
        "prepared_by_position": o.get("author_position", ""),
        "prepared_by_name": o.get("author_name_short", ""),
        "org_short": o.get("short_name", ""),
        "default_approved_by": approvers,
    }


# Backward-compat: scripts in ~/.docs-plugin/ord/scripts/ may import ORG_CONFIGS directly.
class _LazyOrgConfigs(dict):
    """Dict, загружающий конфиг при первом обращении по любому ключу."""
    def __missing__(self, key):
        cfg = _build_org_config()
        self[key] = cfg
        return cfg

ORG_CONFIGS = _LazyOrgConfigs()


DOC_TYPE_CONFIGS = {
    "prikaz": {
        "type_name": "ПРИКАЗ",
        "keyword": "приказываю",
        "doc_name_rp": "приказа",        # родительный: настоящего приказа
        "doc_name_tp": "приказом",       # творительный: с приказом
        "doc_name_vp": "приказ",         # винительный
    },
    "rasporyazhenie": {
        "type_name": "РАСПОРЯЖЕНИЕ",
        "keyword": "обязываю",
        "doc_name_rp": "распоряжения",
        "doc_name_tp": "распоряжением",
        "doc_name_vp": "распоряжение",
    },
    "ukazanie": {
        "type_name": "УКАЗАНИЕ",
        "keyword": "указываю",
        "doc_name_rp": "указания",
        "doc_name_tp": "указанием",
        "doc_name_vp": "указание",
    },
}


# ─── Вспомогательные функции ──────────────────────────────────────────────────

def _set_cell_margins_zero(table):
    """Установить внутренние отступы ячеек таблицы в 0 со всех сторон."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def _remove_table_borders(table):
    """Убрать все границы таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _autofit_contents(table):
    """Установить AutoFit Contents для таблицы.

    Таблица занимает 100% ширины листа, столбцы подстраиваются под содержимое.
    Применять ко всем содержательным таблицам (перечни, листы ознакомления и т.д.).
    Не применять к layout-таблицам (подпись, дата/номер).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Ширина таблицы 100% (5000 = 100% в пятидесятых долях процента)
    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    # tblLayout autofit
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)
    # Ширины столбцов → auto
    for row in tbl.findall(qn("w:tr")):
        for tc in row.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is not None:
                    tcW.set(qn("w:type"), "auto")


def _set_rows_no_break(table):
    """Запретить разрыв строк таблицы между страницами (cantSplit)."""
    for row in table.rows:
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "true")
        trPr.append(cantSplit)


def _set_table_borders(table):
    """Установить тонкие границы для таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _run(para, text, size=14, bold=False, font_name="Times New Roman"):
    """Добавить run к параграфу с настройками шрифта и языком ru-RU.

    Если *text* содержит ``\\n``, каждый перенос превращается в мягкий
    разрыв строки (``<w:br/>``, аналог Shift+Enter в Word).
    """
    text = text.replace("\u2014", "\u2013")  # длинное тире -> среднее
    _warn_slash(text)
    parts = text.split("\n")
    run = None
    for i, part in enumerate(parts):
        if i > 0:
            br = OxmlElement("w:br")
            run._element.append(br)
        if part or i == 0:
            run = para.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold
            rPr = run._element.rPr
            rPr.rFonts.set(qn("w:eastAsia"), font_name)
            lang = rPr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rPr.append(lang)
            lang.set(qn("w:val"), "ru-RU")
            lang.set(qn("w:eastAsia"), "ru-RU")
            lang.set(qn("w:bidi"), "ru-RU")
    return run


def _para(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, space_after=None, space_before=None):
    """Добавить параграф с настройками."""
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    if text:
        _run(p, text, size=size, bold=bold)
    return p


def _dash_para(doc, text, size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Буллет-пункт с «–» (средним тире). Отступ слева 1.25 см, без красной строки."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    _run(p, f"– {text}", size=size)
    return p


def _cell_text(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               vertical_align="center"):
    """Установить текст ячейки с форматированием.

    vertical_align: "center" | "top" | "bottom" | None.
    По умолчанию "center" – вертикальное выравнивание посередине.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if text:
        _run(para, text, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _cell_lines(cell, lines, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_align="center"):
    """Установить несколько строк в ячейке как отдельные параграфы (без \\n).

    lines: list[str] – каждая строка станет отдельным параграфом.
    """
    cell.text = ""
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        if line:
            _run(para, line, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _setup_ord_numbering(doc):
    """Зарегистрировать определение нумерации ОРД в документе.

    Создаёт abstractNum с уровнями:
    - ilvl=0 → 1., 2., 3., ...
    - ilvl=1 → 1.1., 1.2., ...
    - ilvl=2 → 1.1.1., ...  (всего 9 уровней)

    Возвращает numId (int). Передавай в _list_para().
    Для нового независимого списка (например, в приложении) используй _add_num_instance().
    """
    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _lvl_xml(ilvl):
        lvl_text = ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + "."
        return (
            f'<w:lvl xmlns:w="{W}" w:ilvl="{ilvl}">'
            f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
            f'<w:lvlText w:val="{lvl_text}"/><w:suff w:val="space"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="0" w:firstLine="720"/></w:pPr>'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'<w:sz w:val="28"/>'
            f'</w:rPr>'
            f'</w:lvl>'
        )

    # Получаем или создаём numbering part
    REL_NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument"
        "/2006/relationships/numbering"
    )
    try:
        num_part = doc.part.numbering_part
        root = num_part._element
    except (ValueError, AttributeError):
        # Создаём новый numbering part
        CT_NUMBERING = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.numbering+xml"
        )
        empty_xml = f'<w:numbering xmlns:w="{W}"/>'.encode("utf-8")
        root = etree.fromstring(empty_xml)
        num_part = NumberingPart(
            PackURI("/word/numbering.xml"), CT_NUMBERING, root, doc.part.package
        )
        doc.part.relate_to(num_part, REL_NUMBERING)

    # Определяем следующие свободные ID
    existing_abstract = root.findall(qn("w:abstractNum"))
    abs_id = len(existing_abstract)

    existing_nums = root.findall(qn("w:num"))
    num_id = len(existing_nums) + 1

    # Добавляем abstractNum с уникальными идентификаторами.
    # ВАЖНО: по спецификации OOXML все w:abstractNum должны предшествовать
    # всем w:num. Word в Compatibility Mode при нарушении этого порядка
    # игнорирует startOverride и наследует счётчик предыдущего списка.
    # Поэтому вставляем abstractNum перед первым существующим w:num.
    unique_hex = f"{abs_id:08X}"   # например: "00000009", "0000000A", ...
    abs_xml = (
        f'<w:abstractNum xmlns:w="{W}" w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="{unique_hex}"/>'
        f'<w:multiLevelType w:val="multilevel"/>'
        f'<w:tmpl w:val="{unique_hex}"/>'
        + "".join(_lvl_xml(i) for i in range(9))
        + f'</w:abstractNum>'
    )
    first_num_el = root.find(qn("w:num"))
    if first_num_el is not None:
        first_num_index = list(root).index(first_num_el)
        root.insert(first_num_index, etree.fromstring(abs_xml))
    else:
        root.append(etree.fromstring(abs_xml))

    # Добавляем num в конец (после всех abstractNum и существующих num).
    # startOverride гарантирует сброс счётчика независимо от предыдущих списков.
    lvl_overrides = "".join(
        f'<w:lvlOverride w:ilvl="{i}"><w:startOverride w:val="1"/></w:lvlOverride>'
        for i in range(9)
    )
    num_xml = (
        f'<w:num xmlns:w="{W}" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'{lvl_overrides}'
        f'</w:num>'
    )
    root.append(etree.fromstring(num_xml))

    return num_id


def _add_num_instance(doc):
    """Добавить новый независимый экземпляр нумерации (для приложений).

    Нумерация начинается заново с 1. Требует предварительного вызова
    _setup_ord_numbering(). Возвращает новый numId (int).
    """
    return _setup_ord_numbering(doc)


def _list_para(doc, text, level=0, num_id=1, size=14, bold=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Добавить элемент нумерованного списка (встроенная нумерация Word).

    level: 0 – основной пункт (1., 2., ...), 1 – подпункт (1.1., ...).
    num_id: получен из _setup_ord_numbering() или _add_num_instance().
    Отступ первой строки 1.25 см задаётся через abstractNum (не в pPr параграфа).
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    if text:
        _run(p, text, size=size, bold=bold)

    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(level))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.insert(0, numPr)

    return p


def _roman(n):
    """Преобразовать целое число в строку с римской цифрой (I, II, III, ...)."""
    vals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = ""
    for value, numeral in vals:
        while n >= value:
            result += numeral
            n -= value
    return result


def _setup_polozheniye_numbering(doc):
    """Нумерация для Положения (многоуровневый список Word).

    ilvl=0 → «I.», «II.», ... – заголовки разделов (upperRoman, жирный)
    ilvl=1 → «1.1.», «1.2.», ... – пункты (decimal + isLgl)
    ilvl=2 → «1.1.1.», ...        – подпункты

    isLgl на уровнях 1+ гарантирует, что Roman-счётчик уровня 0
    отображается как арабское число в форматах «1.1.», «2.3.» и т.д.
    Без isLgl было бы «I.1.», «II.3.».

    Возвращает numId (int).
    """
    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument"
        "/2006/relationships/numbering"
    )
    try:
        num_part = doc.part.numbering_part
        root = num_part._element
    except (ValueError, AttributeError):
        CT_NUMBERING = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.numbering+xml"
        )
        empty_xml = f'<w:numbering xmlns:w="{W}"/>'.encode("utf-8")
        root = etree.fromstring(empty_xml)
        num_part = NumberingPart(
            PackURI("/word/numbering.xml"), CT_NUMBERING, root, doc.part.package
        )
        doc.part.relate_to(num_part, REL_NUMBERING)

    existing_abstract = root.findall(qn("w:abstractNum"))
    abs_id = len(existing_abstract)
    existing_nums = root.findall(qn("w:num"))
    num_id = len(existing_nums) + 1
    unique_hex = f"{abs_id:08X}"

    # Уровень 0: Roman (заголовки разделов, жирный)
    lvl0 = (
        f'<w:lvl xmlns:w="{W}" w:ilvl="0">'
        f'<w:start w:val="1"/><w:numFmt w:val="upperRoman"/>'
        f'<w:lvlText w:val="%1."/><w:suff w:val="space"/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr>'
        f'<w:rPr><w:b/>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="28"/></w:rPr>'
        f'</w:lvl>'
    )

    # Уровни 1-8: decimal.decimal... + isLgl (Roman→Arabic для ссылок %1, %2...)
    def _lvl_n(ilvl):
        lvl_text = ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + "."
        return (
            f'<w:lvl xmlns:w="{W}" w:ilvl="{ilvl}">'
            f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
            f'<w:isLgl/>'
            f'<w:lvlText w:val="{lvl_text}"/><w:suff w:val="space"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="0" w:firstLine="720"/></w:pPr>'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'<w:sz w:val="28"/></w:rPr>'
            f'</w:lvl>'
        )

    abs_xml = (
        f'<w:abstractNum xmlns:w="{W}" w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="{unique_hex}"/>'
        f'<w:multiLevelType w:val="multilevel"/>'
        f'<w:tmpl w:val="{unique_hex}"/>'
        + lvl0
        + "".join(_lvl_n(i) for i in range(1, 9))
        + f'</w:abstractNum>'
    )
    first_num_el = root.find(qn("w:num"))
    if first_num_el is not None:
        root.insert(list(root).index(first_num_el), etree.fromstring(abs_xml))
    else:
        root.append(etree.fromstring(abs_xml))

    lvl_overrides = "".join(
        f'<w:lvlOverride w:ilvl="{i}"><w:startOverride w:val="1"/></w:lvlOverride>'
        for i in range(9)
    )
    num_xml = (
        f'<w:num xmlns:w="{W}" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'{lvl_overrides}'
        f'</w:num>'
    )
    root.append(etree.fromstring(num_xml))
    return num_id


def _setup_dash_numbering(doc):
    """Нумерация для маркированных списков с «–» в Положениях.

    ilvl=0 → «–» (bullet, en dash)

    Отступ: первая строка 1.25 см (720 twips), остальной текст 0 см.
    Возвращает numId (int).
    """
    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument"
        "/2006/relationships/numbering"
    )
    try:
        num_part = doc.part.numbering_part
        root = num_part._element
    except (ValueError, AttributeError):
        CT_NUMBERING = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.numbering+xml"
        )
        empty_xml = f'<w:numbering xmlns:w="{W}"/>'.encode("utf-8")
        root = etree.fromstring(empty_xml)
        num_part = NumberingPart(
            PackURI("/word/numbering.xml"), CT_NUMBERING, root, doc.part.package
        )
        doc.part.relate_to(num_part, REL_NUMBERING)

    existing_abstract = root.findall(qn("w:abstractNum"))
    abs_id = len(existing_abstract)
    existing_nums = root.findall(qn("w:num"))
    num_id = len(existing_nums) + 1
    unique_hex = f"{abs_id:08X}"

    lvl = (
        f'<w:lvl xmlns:w="{W}" w:ilvl="0">'
        f'<w:start w:val="1"/>'
        f'<w:numFmt w:val="bullet"/>'
        f'<w:lvlText w:val="\u2013"/>'
        f'<w:suff w:val="space"/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="0" w:firstLine="720"/></w:pPr>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="28"/>'
        f'</w:rPr>'
        f'</w:lvl>'
    )
    abs_xml = (
        f'<w:abstractNum xmlns:w="{W}" w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="{unique_hex}"/>'
        f'<w:multiLevelType w:val="singleLevel"/>'
        f'<w:tmpl w:val="{unique_hex}"/>'
        + lvl
        + f'</w:abstractNum>'
    )
    first_num_el = root.find(qn("w:num"))
    if first_num_el is not None:
        root.insert(list(root).index(first_num_el), etree.fromstring(abs_xml))
    else:
        root.append(etree.fromstring(abs_xml))

    num_xml = (
        f'<w:num xmlns:w="{W}" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'</w:num>'
    )
    root.append(etree.fromstring(num_xml))
    return num_id


def add_polozheniye(doc, sections, size=14):
    """Добавить текст положения в документ.

    Использует встроенную нумерацию Word (многоуровневый список):
    - Заголовки разделов: ilvl=0 → «I.», «II.» (Roman, жирный, по центру)
    - Пункты: ilvl=1 → «1.1.», «1.2.» (первая строка 1.25 см, остальное 0)
    - Подпункты: ilvl=2 → «1.1.1.» и т.д.

    isLgl гарантирует отображение Roman-номеров разделов как арабских
    в форматах «1.1.», «2.3.» (без него было бы «I.1.», «II.3.»).

    Параметры
    ----------
    doc : Document
    sections : list[dict]
        {"title": "Общие положения", "points": ["Пункт 1.1.", ("Подпункт", 2)]}
    size : int
        Размер шрифта (по умолчанию 14).
    """
    pol_num_id = _setup_polozheniye_numbering(doc)
    dash_num_id = _setup_dash_numbering(doc)

    for section in sections:
        title = section.get("title", "")
        points = section.get("points", [])

        # Заголовок раздела через список (ilvl=0 → «I.», «II.», ...)
        _para(doc, "", space_after=0)
        _list_para(doc, title, level=0, num_id=pol_num_id, size=size,
                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(doc, "", space_after=0)

        # Пункты через список (ilvl=depth: 1 → «1.1.», 2 → «1.1.1.»)
        # Поддержка маркированных списков:
        #   {"text": "Вводный текст:", "dash_items": ["элемент;", "последний."]}
        for point_item in points:
            if isinstance(point_item, dict) and "dash_items" in point_item:
                text = point_item.get("text", "")
                dash_items = point_item["dash_items"]
                depth = point_item.get("depth", 1)
                _list_para(doc, text, level=depth, num_id=pol_num_id,
                           size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                for dash_text in dash_items:
                    _list_para(doc, dash_text, level=0, num_id=dash_num_id,
                               size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            elif isinstance(point_item, (list, tuple)) and len(point_item) == 2:
                point_text, depth = point_item
                _list_para(doc, point_text, level=depth, num_id=pol_num_id,
                           size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            else:
                point_text, depth = str(point_item), 1
                _list_para(doc, point_text, level=depth, num_id=pol_num_id,
                           size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    return doc


_CATEGORY_BIGRAMS = {"дорожная карта"}
_CATEGORY_WORDS = {
    "положение", "руководство",
    "программа", "инструкция", "форма", "карта",
    "матрица", "структура", "схема",
    "правила", "требования", "рекомендации",
    "состав", "план", "перечень", "порядок", "список",
    "график", "регламент", "кодекс", "реестр",
}


def _auto_break_title(title: str) -> str:
    """Вставить \\n после слова-категории в заголовке приложения.

    Если заголовок уже содержит \\n или не начинается с известной
    категории, возвращает без изменений.
    """
    if not title or "\n" in title:
        return title
    words = title.split()
    low2 = " ".join(w.lower() for w in words[:2]) if len(words) >= 2 else ""
    if low2 in _CATEGORY_BIGRAMS and len(words) > 2:
        return " ".join(words[:2]) + "\n" + " ".join(words[2:])
    if words[0].lower() in _CATEGORY_WORDS and len(words) > 1:
        return words[0] + "\n" + " ".join(words[1:])
    return title


def _utverzhden_form(title):
    """Определить склонение «УТВЕРЖДЕН» по заголовку приложения.

    Возвращает: "УТВЕРЖДЕН", "УТВЕРЖДЕНО", "УТВЕРЖДЕНА" или "УТВЕРЖДЕНЫ".
    """
    first_word = title.strip().split()[0].lower() if title.strip() else ""
    fem = {"программа", "инструкция", "форма", "карта", "дорожная",
           "матрица", "структура", "схема"}
    neut = {"положение", "руководство"}
    plur = {"правила", "требования", "рекомендации"}
    masc = {"состав", "план", "перечень", "порядок", "список",
            "график", "регламент", "кодекс", "реестр"}
    if first_word in fem:
        return "УТВЕРЖДЕНА"
    if first_word in neut:
        return "УТВЕРЖДЕНО"
    if first_word in plur:
        return "УТВЕРЖДЕНЫ"
    # По умолчанию – мужской род (большинство заголовков)
    return "УТВЕРЖДЕН"


def _add_commission_table(doc, members):
    """Добавить таблицу состава комиссии (без границ, 3 столбца).

    members : list[dict]
        Каждый элемент: {"fio": "Фамилия\\nИмя Отчество", "position": "должность",
                         "role": "председатель" | None}.
    """
    ct = doc.add_table(rows=len(members), cols=3)
    _remove_table_borders(ct)

    W_FIO = 4320   # ~7.5 cm
    W_DASH = 425   # ~0.75 cm (минимум под один символ с отступами)
    W_POS = 4609   # остаток до ~16.5 cm (ширина листа)

    # Fixed layout + явная общая ширина таблицы
    tbl_el = ct._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW_el = OxmlElement("w:tblW")
    tblW_el.set(qn("w:w"), str(W_FIO + W_DASH + W_POS))
    tblW_el.set(qn("w:type"), "dxa")
    tblPr.insert(0, tblW_el)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Перезаписываем tblGrid – в Compatibility Mode Word использует именно его,
    # а не tcW, для определения ширины столбцов при fixed layout
    tblGrid = tbl_el.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl_el.append(tblGrid)
    else:
        for gc in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(gc)
    for w in [W_FIO, W_DASH, W_POS]:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    for row_idx, member in enumerate(members):
        fio_parts = member["fio"].split("\n")
        fio_lines = fio_parts + [""]

        c0 = ct.cell(row_idx, 0)
        _cell_lines(c0, fio_lines, align=WD_ALIGN_PARAGRAPH.LEFT,
                    vertical_align="top")
        tc0 = c0._tc
        tcPr0 = tc0.get_or_add_tcPr()
        for e in tcPr0.findall(qn("w:tcW")):
            tcPr0.remove(e)
        tcW0 = OxmlElement("w:tcW")
        tcW0.set(qn("w:w"), str(W_FIO))
        tcW0.set(qn("w:type"), "dxa")
        tcPr0.insert(0, tcW0)

        c1 = ct.cell(row_idx, 1)
        _cell_text(c1, "–", align=WD_ALIGN_PARAGRAPH.CENTER,
                   vertical_align="top")
        tc1 = c1._tc
        tcPr1 = tc1.get_or_add_tcPr()
        for e in tcPr1.findall(qn("w:tcW")):
            tcPr1.remove(e)
        tcW1 = OxmlElement("w:tcW")
        tcW1.set(qn("w:w"), str(W_DASH))
        tcW1.set(qn("w:type"), "dxa")
        tcPr1.insert(0, tcW1)

        pos_text = member["position"]
        role = (member.get("role") or "").strip()
        # Роль в скобках – только для именованных позиций.
        # Рядовые члены («член комиссии», «член рабочей группы» и т.п.)
        # не получают подписи в скобках.
        _NAMED_ROLES = {"председатель", "заместитель председателя", "секретарь"}
        if role.lower() in _NAMED_ROLES:
            pos_text += f" ({role})"
        pos_lines = [pos_text, ""]

        c2 = ct.cell(row_idx, 2)
        _cell_lines(c2, pos_lines, align=WD_ALIGN_PARAGRAPH.LEFT,
                    vertical_align="top")
        tc2 = c2._tc
        tcPr2 = tc2.get_or_add_tcPr()
        for e in tcPr2.findall(qn("w:tcW")):
            tcPr2.remove(e)
        tcW2 = OxmlElement("w:tcW")
        tcW2.set(qn("w:w"), str(W_POS))
        tcW2.set(qn("w:type"), "dxa")
        tcPr2.insert(0, tcW2)

    return ct


def _add_attachment_block(doc, att, att_number, total_attachments, org_cfg, dt_cfg):
    """Добавить один блок приложения (гриф + содержимое).

    att : dict
        {"title": str, "type": "polozheniye"|"commission"|"form_table"|"text",
         "sections": [...], "members": [...], "content": str,
         "columns": [...], "rows": int|list, "col_widths": [...],
         "landscape": bool}  # landscape=True – альбомный формат раздела
    att_number : int (1-based)
    total_attachments : int
    """
    landscape = att.get("landscape", False)
    _add_section_break(doc, landscape=landscape)

    # Гриф
    g = doc.add_table(rows=1, cols=2)
    _remove_table_borders(g)
    g.columns[0].width = Cm(8)
    g.columns[1].width = Cm(8.5)

    att_title = att.get("title", "")
    utv = _utverzhden_form(att_title)

    if total_attachments == 1:
        label = "Приложение"
    else:
        label = f"Приложение № {att_number}"

    att_lines = [
        label, "",
        utv,
        f"{dt_cfg['doc_name_tp']} {org_cfg['org_short']}",
        "от _____________ № ___________",
    ]
    _cell_text(g.cell(0, 0), "")
    _cell_lines(g.cell(0, 1), att_lines, align=WD_ALIGN_PARAGRAPH.LEFT)

    _para(doc, "", space_after=0)

    if att_title:
        _para(doc, "", space_after=0)
        _para(doc, _auto_break_title(att_title), bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        _para(doc, "", space_after=0)

    att_type = att.get("type", "text")

    if att_type == "polozheniye":
        add_polozheniye(doc, att.get("sections", []))
        # Приложения к Положению (sub_attachments)
        sub_atts = att.get("sub_attachments", [])
        if sub_atts:
            # Название Положения без слова «Положение» для грифа
            if att_title.lower().startswith("положение "):
                pol_suffix = att_title[len("Положение "):]
            else:
                pol_suffix = att_title
            total_sub = len(sub_atts)
            for sub_idx, sub_att in enumerate(sub_atts, start=1):
                sub_landscape = sub_att.get("landscape", False)
                _add_section_break(doc, landscape=sub_landscape)
                # Гриф: «Приложение № N / к Положению / <название>»
                sg = doc.add_table(rows=1, cols=2)
                _remove_table_borders(sg)
                sg.columns[0].width = Cm(8)
                sg.columns[1].width = Cm(8.5)
                sub_label = "Приложение" if total_sub == 1 else f"Приложение № {sub_idx}"
                sub_lines = [sub_label, "к Положению", pol_suffix]
                _cell_text(sg.cell(0, 0), "")
                _cell_lines(sg.cell(0, 1), sub_lines,
                            align=WD_ALIGN_PARAGRAPH.LEFT)
                _para(doc, "", space_after=0)
                sub_title = sub_att.get("title", "")
                if sub_title:
                    _para(doc, "", space_after=0)
                    _para(doc, sub_title, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
                    _para(doc, "", space_after=0)
                sub_type = sub_att.get("type", "text")
                if sub_type == "form_table":
                    columns = sub_att.get("columns", [])
                    rows = sub_att.get("rows", 5)
                    sub_col_widths = sub_att.get("col_widths")
                    if columns:
                        filled = isinstance(rows, list)
                        n_data = len(rows) if filled else rows
                        tbl = doc.add_table(rows=1 + n_data, cols=len(columns))
                        _set_table_borders(tbl)
                        _autofit_contents(tbl)
                        _set_rows_no_break(tbl)
                        if sub_col_widths:
                            for i, w in enumerate(sub_col_widths):
                                tbl.columns[i].width = w
                        for i, h in enumerate(columns):
                            _cell_text(tbl.cell(0, i), h, bold=True,
                                       align=WD_ALIGN_PARAGRAPH.CENTER)
                        if filled:
                            for r_idx, row_data in enumerate(rows):
                                for c_idx, cell_val in enumerate(row_data):
                                    cell = tbl.cell(r_idx + 1, c_idx)
                                    val = str(cell_val)
                                    if "\n" in val:
                                        _cell_lines(cell, val.split("\n"))
                                    else:
                                        _cell_text(cell, val)
                        else:
                            for r in range(1, n_data + 1):
                                for c in range(len(columns)):
                                    _cell_text(tbl.cell(r, c), "")
                elif sub_type == "text":
                    content = sub_att.get("content", "")
                    if content:
                        _para(doc, content,
                              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=1.25)
    elif att_type == "commission":
        _add_commission_table(doc, att.get("members", []))
    elif att_type == "form_table":
        columns = att.get("columns", [])
        rows = att.get("rows", 5)
        col_widths = att.get("col_widths")
        if columns:
            filled = isinstance(rows, list)
            n_data = len(rows) if filled else rows
            tbl = doc.add_table(rows=1 + n_data, cols=len(columns))
            _set_table_borders(tbl)
            _autofit_contents(tbl)
            _set_rows_no_break(tbl)
            if col_widths:
                for i, w in enumerate(col_widths):
                    tbl.columns[i].width = w
            for i, h in enumerate(columns):
                _cell_text(tbl.cell(0, i), h, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            if filled:
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_val in enumerate(row_data):
                        cell = tbl.cell(r_idx + 1, c_idx)
                        val = str(cell_val)
                        if "\n" in val:
                            _cell_lines(cell, val.split("\n"))
                        else:
                            _cell_text(cell, val)
            else:
                for r in range(1, n_data + 1):
                    for c in range(len(columns)):
                        _cell_text(tbl.cell(r, c), "")
    elif att_type == "text":
        content = att.get("content", "")
        if content:
            _para(doc, content, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=1.25)


def _add_page_break(doc):
    """Добавить разрыв страницы (устаревшая, используй _add_section_break)."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_section_break(doc, landscape=False):
    """Добавить разрыв раздела с новой страницы.

    Каждый новый раздел:
    - начинается с новой страницы
    - нумерация страниц с 1
    - первая страница раздела без номера
    - номер страницы по центру вверху (со 2-й страницы раздела)
    - поля наследуются от предыдущего раздела

    landscape : bool
        True – альбомная ориентация (A4 29.7×21 см).
    """
    from docx.enum.section import WD_SECTION_START

    # Добавляем пустой параграф-носитель, чтобы sectPr гарантированно
    # прикрепился к отдельному параграфу, а не к элементу таблицы
    carrier = doc.add_paragraph()
    carrier.paragraph_format.space_after = Pt(0)
    carrier.paragraph_format.space_before = Pt(0)

    # Копируем поля из текущего последнего раздела
    prev = doc.sections[-1]
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = prev.top_margin
    new_section.bottom_margin = prev.bottom_margin
    new_section.left_margin = prev.left_margin
    new_section.right_margin = prev.right_margin

    # Ориентацию всегда задаём явно – не наследуем от предыдущего раздела,
    # чтобы альбомное приложение не «протекало» в следующий раздел.
    from docx.enum.section import WD_ORIENT
    if landscape:
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Cm(29.7)
        new_section.page_height = Cm(21)
    else:
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Cm(21)
        new_section.page_height = Cm(29.7)

    # Явно прописываем w:type = "nextPage" в sectPr параграфа,
    # чтобы Word гарантированно создавал разрыв на новую страницу.
    # add_section() мог создать свой параграф, ищем его в body.
    body = doc.element.body
    for p_el in body.findall(qn("w:p")):
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            sect = pPr.find(qn("w:sectPr"))
            if sect is not None:
                # Проверяем, нет ли уже type="nextPage"
                existing_type = sect.find(qn("w:type"))
                if existing_type is None:
                    sect_type = OxmlElement("w:type")
                    sect_type.set(qn("w:val"), "nextPage")
                    sect.insert(0, sect_type)

    # Включить «Особый колонтитул первой страницы»
    new_section.different_first_page_header_footer = True

    # Перезапуск нумерации с 1
    sectPr = new_section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)

    # Номер страницы в верхнем колонтитуле (со 2-й страницы)
    header = new_section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.space_before = Pt(0)

    # Поле PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run = hp.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

    # Первая страница - пустой верхний колонтитул (номер не отображается)
    first_header = new_section.first_page_header
    first_header.is_linked_to_previous = False

    return new_section


def add_notify_sheet(doc, title, doc_type, notify_persons, org_cfg=None):
    """Добавить лист ознакомления как отдельный раздел (новая страница).

    Вынесено из create_ord(), чтобы скрипты с постобработкой могли вызвать
    его вручную в самом конце – после всех приложений, включая кастомные.

    Параметры
    ----------
    doc : Document
    title : str          Заголовок документа (используется в шапке листа).
    doc_type : str       "prikaz" | "rasporyazhenie" | "ukazanie"
    notify_persons : list[dict]  {"name": str, "position": str}
    org_cfg : dict | None  Если None – загружается через _build_org_config().
    """
    if org_cfg is None:
        org_cfg = _build_org_config()
    dt_cfg = DOC_TYPE_CONFIGS[doc_type]

    _add_section_break(doc)

    doc_name_tp = dt_cfg["doc_name_tp"]
    _para(doc, "Лист ознакомления", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, f"с {doc_name_tp} {org_cfg['org_short']}",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, "от «___» __________20__ г. № ____",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    if title:
        _para(doc, f"«{title}»", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, "", space_after=0)

    notify_persons = sorted(notify_persons, key=lambda p: p.get("name", ""))

    num_rows = max(len(notify_persons), 13)
    ozn_table = doc.add_table(rows=1 + num_rows, cols=4)
    _set_table_borders(ozn_table)
    # Равная ширина столбцов (distribute columns): 100% / 4 = 25% каждый
    tbl = ozn_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    col_w_pct = "1250"  # 5000 / 4
    for row in tbl.findall(qn("w:tr")):
        for tc in row.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), col_w_pct)
            tcW.set(qn("w:type"), "pct")

    for i, h in enumerate(["Наименование должности", "Подпись",
                            "Расшифровка подписи", "Дата"]):
        _cell_text(ozn_table.cell(0, i), h, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, person in enumerate(notify_persons):
        row = idx + 1
        _cell_text(ozn_table.cell(row, 0), person.get("position", ""))
        _cell_text(ozn_table.cell(row, 1), "")
        _cell_text(ozn_table.cell(row, 2), person.get("name", ""))
        _cell_text(ozn_table.cell(row, 3), "")

    for idx in range(len(notify_persons), num_rows):
        for col in range(4):
            _cell_text(ozn_table.cell(idx + 1, col), "")


# ─── Основная функция ─────────────────────────────────────────────────────────

def create_ord(
    org=None,  # deprecated, ignored; kept for backward compat with scripts
    doc_type="prikaz",
    title="",
    basis="",
    points=None,
    control_person="",
    notify_persons=None,
    approved_by=None,
    has_attachment=False,
    attachment_title="",
    attachment_content="",
    attachment_polozheniye=None,
    attachments=None,
    doc_date=None,
    doc_number="",
    output_path=None,
):
    """
    Создаёт организационно-распорядительный документ в формате .docx.

    Параметры
    ----------
    doc_type : str
        Тип документа: "prikaz" | "rasporyazhenie" | "ukazanie"
    title : str
        Заголовок документа (например, "О назначении ответственных за ...")
    basis : str
        Обоснование – текст перед ключевым словом. Не включай само ключевое
        слово ("приказываю" и т.д.) – оно добавляется автоматически.
    points : list[str | tuple[str, int]]
        Пункты документа. Нумерация добавляется автоматически (Word numbering).
        Для основных пунктов передавай строку: "Текст пункта".
        Для подпунктов передавай кортеж: ("Текст подпункта", level),
        где level=0 – основной (1., 2.), level=1 – подпункт (1.1., 2.1.).
        НЕ добавляй номера вручную – они генерируются автоматически.
    control_person : str
        На кого возложить контроль (в винительном падеже, без "на").
        Пример: "заместителя генерального директора по стратегическому развитию [Фамилия И.О.]"
        Для варианта «оставляю за собой» передать "self".
    notify_persons : list[dict]
        Лист ознакомления. Каждый элемент: {"name": "ФИО", "position": "Должность"}.
    approved_by : list[dict]
        Согласующие. Формат: {"name": "ФИО", "position": "Должность"}.
    has_attachment : bool
        Включать ли страницу приложения (для обратной совместимости).
    attachment_title : str
        Заголовок приложения (для обратной совместимости).
    attachment_content : str
        Содержимое приложения (для обратной совместимости).
    attachment_polozheniye : list[dict] | None
        Положение как приложение (для обратной совместимости).
    attachments : list[dict] | None
        Множественные приложения. Имеет приоритет над has_attachment и др.
        Каждый элемент – словарь:
        - {"title": str, "type": "polozheniye", "sections": list[dict]}
        - {"title": str, "type": "commission", "members": list[dict]}
          members: [{"fio": "Фамилия\\nИмя Отчество", "position": str,
                     "role": str|None}]
        - {"title": str, "type": "form_table", "columns": list[str],
           "rows": int | list[list[str]],
           "col_widths": list[Cm] | None}
          rows=int → пустая таблица с N строками (по умолчанию 5).
          rows=list → заполненная таблица; если ячейка содержит \\n,
          она разбивается на несколько строк внутри ячейки.
          col_widths – опционально, ширины столбцов (Cm).
        - {"title": str, "type": "text", "content": str}
        Нумерация «Приложение № N» и склонение «УТВЕРЖДЕН/О/А/Ы»
        определяются автоматически.
    doc_date : str | None
        Дата документа "ДД.ММ.ГГГГ". По умолчанию – сегодня.
    doc_number : str
        Номер документа (если известен).
    output_path : str | None
        Путь для сохранения. По умолчанию – "<title>.docx".
    """
    if points is None:
        points = []
    if notify_persons is None:
        notify_persons = []
    if approved_by is None:
        approved_by = []
    # Конвертация старых параметров приложений в новый формат attachments
    if attachments is None:
        attachments = []
        if attachment_polozheniye:
            attachments.append({
                "title": attachment_title,
                "type": "polozheniye",
                "sections": attachment_polozheniye,
            })
        elif has_attachment and attachment_content:
            attachments.append({
                "title": attachment_title,
                "type": "text",
                "content": attachment_content,
            })
        elif has_attachment and attachment_title:
            attachments.append({
                "title": attachment_title,
                "type": "text",
                "content": "",
            })
    _date_was_provided = doc_date is not None
    if doc_date is None:
        doc_date = date.today().strftime("%d.%m.%Y")

    org_cfg = _build_org_config()
    dt_cfg = DOC_TYPE_CONFIGS[doc_type]

    # Если approved_by не передан – использовать дефолтных согласующих организации
    if not approved_by and "default_approved_by" in org_cfg:
        approved_by = [dict(p) for p in org_cfg["default_approved_by"]]

    if output_path is None:
        safe_title = title.replace("/", "_").replace("\\", "_")[:80] if title else dt_cfg["type_name"]
        output_path = f"{safe_title}.docx"

    doc = Document()

    # ── Поля страницы (A4) ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    # Первый раздел: особый колонтитул первой страницы, нумерация с 1
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)

    # Основной header (со 2-й страницы) - номер по центру вверху
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.space_before = Pt(0)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run0 = hp.add_run()
    run0.font.name = "Times New Roman"
    run0.font.size = Pt(12)
    run0._element.append(fldChar1)
    run0._element.append(instrText)
    run0._element.append(fldChar2)

    # Первая страница - пустой верхний колонтитул
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False

    # Шрифт и язык по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    # Язык документа – русский (чтобы Word не помечал текст как English)
    style_rPr = style.element.get_or_add_rPr()
    style_lang = style_rPr.find(qn("w:lang"))
    if style_lang is None:
        style_lang = OxmlElement("w:lang")
        style_rPr.append(style_lang)
    style_lang.set(qn("w:val"), "ru-RU")
    style_lang.set(qn("w:eastAsia"), "ru-RU")
    style_lang.set(qn("w:bidi"), "ru-RU")
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    # Нумерация пунктов (встроенные списки Word)
    doc_num_id = _setup_ord_numbering(doc)

    # ── Шапка организации ─────────────────────────────────────────────────
    header_size = org_cfg.get("header_font_size", 14)
    for line_item in org_cfg["header_lines"]:
        # Формат: (text, bold, italic, bottom_border, font_size_override)
        line_text = line_item[0]
        line_bold = line_item[1]
        line_italic = line_item[2] if len(line_item) > 2 else False
        line_border = line_item[3] if len(line_item) > 3 else False
        line_size = line_item[4] if len(line_item) > 4 else header_size

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if line_text:
            r = _run(p, line_text, size=line_size, bold=line_bold)
            if line_italic:
                r.italic = True

        if line_border:
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)

    # Одна пустая строка после шапки
    _para(doc, "", space_after=0)

    # ── Тип документа ─────────────────────────────────────────────────────
    _para(doc, dt_cfg["type_name"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=0)
    _para(doc, "", space_after=0)

    # ── Дата и номер ──────────────────────────────────────────────────────
    date_num_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(date_num_table)
    date_num_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    date_num_table.columns[0].width = Cm(8)
    date_num_table.columns[1].width = Cm(8.5)

    date_str = f"от «{doc_date[0:2]}» {_month_name(doc_date)} {doc_date[6:10]} г." if _date_was_provided else "от «____» __________ 20__ г."
    num_str = f"№ {doc_number}" if doc_number else "№ ______"
    _cell_text(date_num_table.cell(0, 0), date_str, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(date_num_table.cell(0, 1), num_str, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Пустая строка
    _para(doc, "", space_after=0)

    # ── Заголовок ─────────────────────────────────────────────────────────
    if title:
        # Таблица 1x2, без границ, cell margins 0, текст обычный
        title_table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(title_table)
        _set_cell_margins_zero(title_table)
        _cell_text(title_table.cell(0, 0), title, bold=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text(title_table.cell(0, 1), "",
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _para(doc, "", space_after=0)

    # ── Обоснование + ключевое слово ─────────────────────────────────────
    if basis:
        keyword = dt_cfg["keyword"]
        if org_cfg["keyword_upper"]:
            keyword = keyword.upper()

        basis_clean = basis.rstrip(" ,.")

        # Обоснование отдельно, пустая строка, ПРИКАЗЫВАЮ: жирный, пустая строка
        # Запятая после обоснования НЕ ставится.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        _run(p, basis_clean)

        _para(doc, "", space_after=0)

        p_kw = doc.add_paragraph()
        p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_kw.paragraph_format.first_line_indent = Cm(0)
        p_kw.paragraph_format.line_spacing = 1.0
        p_kw.paragraph_format.space_after = Pt(0)
        _run(p_kw, f"{keyword}:", bold=True)

        _para(doc, "", space_after=0)

    # ── Пункты документа ─────────────────────────────────────────────────
    # points: list[str | (str, int)]
    #   str           → уровень 0 (основной пункт)
    #   (str, level)  → явно заданный уровень (0, 1, 2, ...)
    for point_item in points:
        if isinstance(point_item, (list, tuple)) and len(point_item) == 2:
            point_text, level = point_item
        else:
            point_text, level = str(point_item), 0
        _list_para(doc, point_text, level=level, num_id=doc_num_id)

    # ── Пункт о контроле ─────────────────────────────────────────────────
    if control_person:
        doc_name_rp = dt_cfg["doc_name_rp"]
        if control_person.lower() in ("self", "за собой", "оставляю за собой"):
            _list_para(
                doc,
                f"Контроль за исполнением настоящего {doc_name_rp} оставляю за собой.",
                level=0,
                num_id=doc_num_id,
            )
        else:
            cp = control_person.rstrip(".")
            _list_para(
                doc,
                f"Контроль за исполнением настоящего {doc_name_rp} возложить на {cp}.",
                level=0,
                num_id=doc_num_id,
            )

    # Три пустые строки перед подписью
    _para(doc, "", space_after=0)
    _para(doc, "", space_after=0)
    _para(doc, "", space_after=0)

    # ── Блок подписи ──────────────────────────────────────────────────────
    sig_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(sig_table)
    _set_cell_margins_zero(sig_table)
    sig_table.columns[0].width = Cm(8.5)
    sig_table.columns[1].width = Cm(8)

    _cell_text(sig_table.cell(0, 0), org_cfg["leader_title"],
               align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(sig_table.cell(0, 1), org_cfg["leader_name"],
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Разрыв раздела: тело приказа → блок согласования ────────────────
    _add_section_break(doc)

    # ── Подготовил + Согласован ────────────────────────────────────────────
    # Единая таблица 2 колонки без границ, cell margins 0
    type_name_lower = dt_cfg["type_name"].lower()
    num_rows = 1 + 1 + 1 + len(approved_by)
    appr_combined = doc.add_table(rows=num_rows, cols=2)
    _remove_table_borders(appr_combined)
    _set_cell_margins_zero(appr_combined)

    # Строка 0: «Проект приказа представил:»
    _cell_lines(appr_combined.cell(0, 0),
                [f"Проект {type_name_lower}а представил:", ""],
                vertical_align=None)
    _cell_text(appr_combined.cell(0, 1), "",
               align=WD_ALIGN_PARAGRAPH.RIGHT, vertical_align=None)

    # Строка 1: автор (должность + имя)
    _cell_lines(appr_combined.cell(1, 0),
                [org_cfg["prepared_by_position"], ""],
                vertical_align=None)
    _cell_text(appr_combined.cell(1, 1), org_cfg["prepared_by_name"],
               align=WD_ALIGN_PARAGRAPH.RIGHT, vertical_align=None)

    # Строка 2: «Проект приказа согласовали:»
    _cell_lines(appr_combined.cell(2, 0),
                ["", f"Проект {type_name_lower}а согласовали:", ""],
                vertical_align=None)
    _cell_text(appr_combined.cell(2, 1), "",
               align=WD_ALIGN_PARAGRAPH.RIGHT, vertical_align=None)

    # Строки 3+: согласующие
    for idx, person in enumerate(approved_by):
        row = 3 + idx
        _cell_lines(appr_combined.cell(row, 0),
                    [person.get("position", ""), ""],
                    vertical_align=None)
        _cell_text(appr_combined.cell(row, 1), person.get("name", ""),
                   align=WD_ALIGN_PARAGRAPH.RIGHT, vertical_align=None)

    # ── Приложения ──────────────────────────────────────────────────────
    total_att = len(attachments)
    for att_idx, att in enumerate(attachments, start=1):
        _add_attachment_block(doc, att, att_idx, total_att, org_cfg, dt_cfg)

    # ── Лист ознакомления ─────────────────────────────────────────────────
    add_notify_sheet(doc, title=title, doc_type=doc_type,
                     notify_persons=notify_persons, org_cfg=org_cfg)

    # ── Сохранение ────────────────────────────────────────────────────────
    output_path = os.path.expanduser(output_path)
    doc.save(output_path)
    print(f"Сохранено: {output_path}")
    log_generation(doc_type, title, output_path, {
        "points_count": len(points),
        "attachments_count": len(attachments),
        "has_notify": bool(notify_persons),
        "has_approval": bool(approved_by),
    })
    return output_path


def _month_name(date_str):
    """Получить название месяца из даты ДД.ММ.ГГГГ."""
    months = {
        "01": "января", "02": "февраля", "03": "марта",
        "04": "апреля", "05": "мая", "06": "июня",
        "07": "июля", "08": "августа", "09": "сентября",
        "10": "октября", "11": "ноября", "12": "декабря",
    }
    month = date_str[3:5]
    return months.get(month, "__________")


if __name__ == "__main__":
    create_ord(
        doc_type="prikaz",
        title="О назначении ответственных за ведение документации",
        basis="В целях повышения эффективности документооборота",
        points=[
            "Назначить ответственных лиц за ведение документации согласно приложению.",
        ],
        control_person="заместителя генерального директора по стратегическому развитию",
        notify_persons=[],
        doc_date=date.today().strftime("%d.%m.%Y"),
        output_path="/tmp/пример_приказ.docx",
    )
    print("Пример приказа создан: /tmp/пример_приказ.docx")
