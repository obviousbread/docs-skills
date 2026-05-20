#!/usr/bin/env python3
"""Генератор должностных инструкций (ДИ).

Конфигурация организации загружается из ~/.docs-plugin/org_details.md через
общий lib/db.py. При отсутствии файла запустите docs-init для первичной
настройки.

Использование:
    python3 generate.py  # создает /tmp/ДИ smoke-test.docx с примером

В скилле:
    from generate import create_di
    create_di(
        position_full="...",
        position_full_genitive="...",
        sections=[...],
        output_path="...",
        approvers=[{"position": "...", "name": "..."}],
    )
"""

import os
import re
import sys as _sys
import warnings
from datetime import date

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_HERE = os.path.dirname(os.path.abspath(__file__))
for _lib_path in (
    os.path.join(_HERE, "..", "..", "lib"),
    os.path.expanduser("~/.docs-plugin/runtime/lib"),
):
    _lib_path = os.path.abspath(_lib_path)
    if os.path.isdir(_lib_path) and _lib_path not in _sys.path:
        _sys.path.insert(0, _lib_path)
from db import log_generation, ORG_DETAILS_PATH


# ─── Метаданные документа ──────────────────────────────────────────────────────

def _new_document():
    """Создать Document и очистить служебные метаданные python-docx."""
    doc = Document()
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.category = ""
    return doc


# ─── Валидация текста ──────────────────────────────────────────────────────────

_SLASH_OK = re.compile(r'п/п|ИНН/КПП|[0-9]/[0-9]|[A-Za-z]/[A-Za-z]')


def _sanitize_text(text):
    """Безопасные автозамены, применяемые ко всему тексту документа.

    - U+2014 («—», длинное тире) → U+2013 («–», среднее тире)
    - «ё» → «е», «Ё» → «Е»
    """
    return text.replace("—", "–").replace("ё", "е").replace("Ё", "Е")


def _warn_slash(text):
    """Предупреждение при «/» между кириллическими словами."""
    for m in re.finditer(r'[а-яА-ЯёЁ]+\s*/\s*[а-яА-ЯёЁ]+', text):
        if not _SLASH_OK.search(m.group()):
            warnings.warn(
                f"Косая черта в тексте: «{m.group()}». "
                f"Используй запятую, «или», «и» или скобки.",
                stacklevel=3,
            )


# ─── Загрузка реквизитов ──────────────────────────────────────────────────────

def _load_org_details():
    """Загрузить реквизиты из ~/.docs-plugin/org_details.md.

    Парсит формат `key: value`. Добавляет алиасы под имена полей, которые
    использует генератор ДИ (approver_*, default_author_*, short_name_genitive).
    """
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        print("WARNING: ~/.docs-plugin/org_details.md not found. Run docs-init to configure.")
        return config
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()

    # Алиасы для блоков ДИ.
    config.setdefault("approver_position", config.get("leader_title", ""))
    config.setdefault("approver_org", config.get("short_name", ""))
    config.setdefault("approver_name", config.get("leader_name_nom", ""))
    config.setdefault("default_author_position", config.get("author_position", ""))
    config.setdefault("default_author_name", config.get("author_name_short", ""))
    config.setdefault("short_name_genitive", config.get("short_name", ""))
    return config


# ─── Базовые хелперы шрифта и таблиц ───────────────────────────────────────────

def _run(para, text, size=14, bold=False, italic=False, font_name="Times New Roman"):
    text = _sanitize_text(text)
    _warn_slash(text)
    parts = text.split("\n")
    run = None
    for i, part in enumerate(parts):
        if i > 0:
            br = OxmlElement("w:br")
            run._element.append(br)
        if part or i == 0:
            run = para.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            rPr = run._element.rPr
            rPr.rFonts.set(qn("w:eastAsia"), font_name)
            lang = rPr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rPr.append(lang)
            lang.set(qn("w:val"), "ru-RU")
            lang.set(qn("w:eastAsia"), "ru-RU")
            lang.set(qn("w:bidi"), "ru-RU")
    return run


def _para(doc, text, size=14, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    if text:
        _run(p, text, size=size, bold=bold, italic=italic)
    return p


def _cell_text(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               vertical_align="center"):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if text:
        _run(para, text, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _cell_lines(cell, lines, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_align="center"):
    cell.text = ""
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        if line:
            _run(para, line, size=size, bold=bold)
    if vertical_align:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical_align)
        tcPr.append(vAlign)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_cell_margins_zero(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    cellMar = OxmlElement("w:tblCellMar")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_rows_no_break(table):
    for row in table.rows:
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "true")
        trPr.append(cantSplit)


def _table_full_width_fixed(table, col_widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(col_widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.insert(0, tblW)
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for row in tbl.findall(qn("w:tr")):
        for tc, w in zip(row.findall(qn("w:tc")), col_widths_dxa):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")


# ─── Нумерация ──────────────────────────────────────────────────────────────────

def _setup_di_numbering(doc):
    """Многоуровневая арабская нумерация. Возвращает numId."""
    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument"
        "/2006/relationships/numbering"
    )

    try:
        num_part = doc.part.numbering_part
        root = num_part._element
    except (ValueError, AttributeError):
        CT_NUMBERING = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.numbering+xml"
        )
        empty_xml = f'<w:numbering xmlns:w="{W}"/>'.encode("utf-8")
        root = etree.fromstring(empty_xml)
        num_part = NumberingPart(
            PackURI("/word/numbering.xml"), CT_NUMBERING, root, doc.part.package
        )
        doc.part.relate_to(num_part, REL_NUMBERING)

    existing_abstract = root.findall(qn("w:abstractNum"))
    abs_id = len(existing_abstract)
    existing_nums = root.findall(qn("w:num"))
    num_id = len(existing_nums) + 1
    unique_hex = f"{abs_id:08X}"

    def _lvl_xml(ilvl):
        lvl_text = ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + "."
        if ilvl == 0:
            ind_xml = '<w:ind w:left="0" w:firstLine="0"/>'
            rpr = (
                '<w:rPr><w:b/>'
                '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                '<w:sz w:val="28"/></w:rPr>'
            )
        else:
            ind_xml = '<w:ind w:left="0" w:firstLine="720"/>'
            rpr = (
                '<w:rPr>'
                '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
                '<w:sz w:val="28"/></w:rPr>'
            )
        return (
            f'<w:lvl xmlns:w="{W}" w:ilvl="{ilvl}">'
            f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
            f'<w:lvlText w:val="{lvl_text}"/><w:suff w:val="space"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr>{ind_xml}</w:pPr>'
            f'{rpr}'
            f'</w:lvl>'
        )

    abs_xml = (
        f'<w:abstractNum xmlns:w="{W}" w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="{unique_hex}"/>'
        f'<w:multiLevelType w:val="multilevel"/>'
        f'<w:tmpl w:val="{unique_hex}"/>'
        + "".join(_lvl_xml(i) for i in range(9))
        + f'</w:abstractNum>'
    )
    first_num_el = root.find(qn("w:num"))
    if first_num_el is not None:
        root.insert(list(root).index(first_num_el), etree.fromstring(abs_xml))
    else:
        root.append(etree.fromstring(abs_xml))

    lvl_overrides = "".join(
        f'<w:lvlOverride w:ilvl="{i}"><w:startOverride w:val="1"/></w:lvlOverride>'
        for i in range(9)
    )
    num_xml = (
        f'<w:num xmlns:w="{W}" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'{lvl_overrides}'
        f'</w:num>'
    )
    root.append(etree.fromstring(num_xml))
    return num_id


def _setup_dash_numbering(doc):
    """Маркированный список с «–». Возвращает numId."""
    from docx.parts.numbering import NumberingPart
    from docx.opc.packuri import PackURI
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_NUMBERING = (
        "http://schemas.openxmlformats.org/officeDocument"
        "/2006/relationships/numbering"
    )
    try:
        num_part = doc.part.numbering_part
        root = num_part._element
    except (ValueError, AttributeError):
        CT_NUMBERING = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.numbering+xml"
        )
        empty_xml = f'<w:numbering xmlns:w="{W}"/>'.encode("utf-8")
        root = etree.fromstring(empty_xml)
        num_part = NumberingPart(
            PackURI("/word/numbering.xml"), CT_NUMBERING, root, doc.part.package
        )
        doc.part.relate_to(num_part, REL_NUMBERING)

    existing_abstract = root.findall(qn("w:abstractNum"))
    abs_id = len(existing_abstract)
    existing_nums = root.findall(qn("w:num"))
    num_id = len(existing_nums) + 1
    unique_hex = f"{abs_id:08X}"

    lvl = (
        f'<w:lvl xmlns:w="{W}" w:ilvl="0">'
        f'<w:start w:val="1"/>'
        f'<w:numFmt w:val="bullet"/>'
        f'<w:lvlText w:val="–"/>'
        f'<w:suff w:val="space"/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="0" w:firstLine="720"/></w:pPr>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="28"/>'
        f'</w:rPr>'
        f'</w:lvl>'
    )
    abs_xml = (
        f'<w:abstractNum xmlns:w="{W}" w:abstractNumId="{abs_id}">'
        f'<w:nsid w:val="{unique_hex}"/>'
        f'<w:multiLevelType w:val="singleLevel"/>'
        f'<w:tmpl w:val="{unique_hex}"/>'
        + lvl
        + f'</w:abstractNum>'
    )
    first_num_el = root.find(qn("w:num"))
    if first_num_el is not None:
        root.insert(list(root).index(first_num_el), etree.fromstring(abs_xml))
    else:
        root.append(etree.fromstring(abs_xml))

    num_xml = (
        f'<w:num xmlns:w="{W}" w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abs_id}"/>'
        f'</w:num>'
    )
    root.append(etree.fromstring(num_xml))
    return num_id


def _list_para(doc, text, level=0, num_id=1, size=14, bold=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _run(p, text, size=size, bold=bold)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(level))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.insert(0, numPr)
    return p


def _add_section_break(doc):
    from docx.enum.section import WD_SECTION_START, WD_ORIENT

    carrier = doc.add_paragraph()
    carrier.paragraph_format.space_after = Pt(0)
    carrier.paragraph_format.space_before = Pt(0)

    prev = doc.sections[-1]
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = prev.top_margin
    new_section.bottom_margin = prev.bottom_margin
    new_section.left_margin = prev.left_margin
    new_section.right_margin = prev.right_margin
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Cm(21)
    new_section.page_height = Cm(29.7)

    body = doc.element.body
    for p_el in body.findall(qn("w:p")):
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            sect = pPr.find(qn("w:sectPr"))
            if sect is not None:
                existing_type = sect.find(qn("w:type"))
                if existing_type is None:
                    sect_type = OxmlElement("w:type")
                    sect_type.set(qn("w:val"), "nextPage")
                    sect.insert(0, sect_type)

    return new_section


# ─── Высокоуровневые блоки ─────────────────────────────────────────────────────

def _add_grif(doc, org):
    grif_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(grif_table)
    _set_cell_margins_zero(grif_table)
    grif_table.columns[0].width = Cm(8.5)
    grif_table.columns[1].width = Cm(8.0)

    _cell_text(grif_table.cell(0, 0), "", align=WD_ALIGN_PARAGRAPH.LEFT)

    right = grif_table.cell(0, 1)
    _cell_lines(right, [
        "«УТВЕРЖДАЮ»",
        org.get("approver_position", ""),
        org.get("approver_org", ""),
        f"________________{org.get('approver_name', '')}",
        "«_____»________________202 ___ года",
    ], align=WD_ALIGN_PARAGRAPH.LEFT, vertical_align="top")

    _para(doc, "", space_after=0)


def _add_title(doc, position_full):
    _para(doc, "ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ", bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, position_full, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para(doc, "", space_after=0)


def _add_section(doc, section, di_num_id, dash_num_id, size=14):
    """Добавить один раздел ДИ.

    section : {"title": str, "intro": str | None,
               "points": list of (str | dict | (str, depth))}.
    """
    title = section.get("title", "")
    intro = section.get("intro")
    points = section.get("points", [])

    _para(doc, "", space_after=0)
    _list_para(doc, title, level=0, num_id=di_num_id, size=size,
               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "", space_after=0)

    if intro:
        _para(doc, intro, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=1.25,
              size=size, space_after=0)

    for item in points:
        if isinstance(item, dict):
            text = item.get("text", "")
            dash_items = item.get("dash_items", [])
            depth = item.get("depth", 1)
            if text:
                _list_para(doc, text, level=depth, num_id=di_num_id,
                           size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            for dash_text in dash_items:
                _list_para(doc, dash_text, level=0, num_id=dash_num_id,
                           size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        elif isinstance(item, (list, tuple)) and len(item) == 2:
            text, depth = item
            _list_para(doc, str(text), level=depth, num_id=di_num_id,
                       size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        else:
            _list_para(doc, str(item), level=1, num_id=di_num_id,
                       size=size, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _add_signoff_table(doc, label, persons):
    _para(doc, label, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)

    n = len(persons)
    if n == 0:
        return

    table = doc.add_table(rows=1 + n, cols=4)
    _set_table_borders(table)
    _set_rows_no_break(table)
    _table_full_width_fixed(table, [3402, 2268, 1984, 1701])

    headers = ["Должность", "ФИО", "Подпись", "Дата"]
    for i, h in enumerate(headers):
        _cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, person in enumerate(persons):
        row = idx + 1
        position = person.get("position", "")
        name = person.get("name", "")

        position_lines = position.split("\n") if "\n" in position else [position]
        _cell_lines(table.cell(row, 0), position_lines + [""])
        _cell_lines(table.cell(row, 1), [name, ""])
        _cell_lines(table.cell(row, 2), ["", ""])
        _cell_lines(table.cell(row, 3), ["", ""])

    _para(doc, "", space_after=0)


def _add_acquaintance_sheet(doc, position_full_genitive, org_short_genitive,
                            num_rows=23):
    _add_section_break(doc)

    _para(doc, "Лист ознакомления", align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=0)
    _para(doc,
          f"С должностной инструкцией {position_full_genitive} "
          f"{org_short_genitive}, ознакомлен(а), один экземпляр получил(а) "
          f"на руки и обязуюсь выполнять ее и хранить на рабочем месте:",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=1.25, space_after=0)
    _para(doc, "", space_after=0)

    table = doc.add_table(rows=1 + num_rows, cols=3)
    _set_table_borders(table)
    _set_rows_no_break(table)
    _table_full_width_fixed(table, [3969, 2835, 2552])

    for i, h in enumerate(["ФИО", "Подпись", "Дата"]):
        _cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r in range(1, 1 + num_rows):
        for c in range(3):
            _cell_text(table.cell(r, c), "")


def _resolve_output_path(path):
    """Если файл существует — добавить суффикс ' 2', ' 3', ... перед .docx."""
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    n = 2
    while os.path.exists(f"{base} {n}{ext}"):
        n += 1
    return f"{base} {n}{ext}"


def _default_output_path(org, position_full):
    """Сформировать путь по умолчанию: <output_dir_di>/ДИ <первое слово>.docx."""
    output_dir = org.get("output_dir_di", "")
    head = (position_full or "должность").split(",")[0].strip()
    head = head.split(" ")[0] if head else "должность"
    filename = f"ДИ {head}.docx"
    if output_dir:
        return os.path.join(os.path.expanduser(output_dir), filename)
    return filename


# ─── Главная функция ───────────────────────────────────────────────────────────

def create_di(
    position_full,
    position_full_genitive,
    sections,
    output_path=None,
    author=None,
    approvers=None,
    org=None,
    acquaintance_rows=23,
):
    """Создает должностную инструкцию в формате .docx.

    Параметры
    ----------
    position_full : str
        Полное наименование должности и учреждения для подзаголовка.
    position_full_genitive : str
        Полное наименование должности в родительном падеже для шапки Листа
        ознакомления (без названия учреждения).
    sections : list[dict]
        Шесть разделов ДИ. Каждый: {"title": str, "intro": str|None,
        "points": list of (str | dict | (str, depth))}.
    output_path : str | None
        Путь сохранения .docx. Если None — формируется как
        <output_dir_di>/ДИ <первое слово должности>.docx; output_dir_di берется
        из ~/.docs-plugin/org_details.md. Если файл существует — добавится
        суффикс ` 2`.
    author : dict | None
        {"position": str, "name": str}. По умолчанию — default_author из
        ~/.docs-plugin/org_details.md (author_position / author_name_short).
    approvers : list[dict] | None
        [{"position": str, "name": str}, ...] для блока «СОГЛАСОВАНО».
    org : dict | None
        Реквизиты организации; если None — _load_org_details().
    acquaintance_rows : int
        Количество пустых строк в Листе ознакомления (по умолчанию 23).

    Возвращает фактический путь сохранения.
    """
    if org is None:
        org = _load_org_details()
    if author is None:
        author = {
            "position": org.get("default_author_position", ""),
            "name": org.get("default_author_name", ""),
        }
    if approvers is None:
        approvers = []

    org_short_genitive = org.get("short_name_genitive", org.get("short_name", ""))

    if output_path is None:
        output_path = _default_output_path(org, position_full)
    output_path = os.path.expanduser(output_path)
    final_path = _resolve_output_path(output_path)

    doc = _new_document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style_rPr = style.element.get_or_add_rPr()
    style_lang = style_rPr.find(qn("w:lang"))
    if style_lang is None:
        style_lang = OxmlElement("w:lang")
        style_rPr.append(style_lang)
    style_lang.set(qn("w:val"), "ru-RU")
    style_lang.set(qn("w:eastAsia"), "ru-RU")
    style_lang.set(qn("w:bidi"), "ru-RU")
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    di_num_id = _setup_di_numbering(doc)
    dash_num_id = _setup_dash_numbering(doc)

    _add_grif(doc, org)
    _add_title(doc, position_full)

    for sec in sections:
        _add_section(doc, sec, di_num_id, dash_num_id)

    _para(doc, "", space_after=0)
    _add_signoff_table(doc, "ПРОЕКТ РАЗРАБОТАН:", [author])
    _add_signoff_table(doc, "СОГЛАСОВАНО:", approvers)

    _add_acquaintance_sheet(doc, position_full_genitive, org_short_genitive,
                            num_rows=acquaintance_rows)

    doc.save(final_path)
    print(f"Сохранено: {final_path}")
    log_generation("di", position_full, final_path, {
        "approvers_count": len(approvers),
        "sections_count": len(sections),
    })
    return final_path


# ─── Пример для самотестирования ───────────────────────────────────────────────

if __name__ == "__main__":
    sample_sections = [
        {
            "title": "Общие положения",
            "points": [
                "Настоящая должностная инструкция определяет функциональные "
                "обязанности, права и ответственность специалиста тестового "
                "отдела (далее – Специалист, Отдел, Учреждение).",
                {
                    "text": "На должность Специалиста назначается лицо, имеющее:",
                    "dash_items": [
                        "высшее образование – специалитет или магистратура;",
                        "дополнительное профессиональное образование (при наличии);",
                        "опыт работы по направлению деятельности не менее 1 года.",
                    ],
                },
                "Специалист в своей работе подчиняется непосредственно "
                "начальнику Отдела.",
            ],
        },
        {
            "title": "Трудовая функция",
            "intro": "Специалист осуществляет следующие функции:",
            "points": [
                "Выполнение поручений руководителя в рамках компетенции Отдела.",
                "Подготовка проектов служебных документов по направлениям "
                "деятельности Отдела.",
            ],
        },
        {
            "title": "Должностные обязанности",
            "intro": "Специалист выполняет следующие должностные обязанности:",
            "points": [
                "Выполняет поручения непосредственного руководителя в "
                "установленные сроки.",
            ],
        },
        {
            "title": "Права",
            "intro": "Специалист имеет право:",
            "points": [
                "Запрашивать и получать информационные материалы и нормативные "
                "правовые акты, необходимые для исполнения обязанностей.",
                "Специалист пользуется всеми правами в соответствии с Трудовым "
                "кодексом Российской Федерации.",
            ],
        },
        {
            "title": "Ответственность",
            "intro": "Специалист несет ответственность за:",
            "points": [
                "неисполнение или ненадлежащее исполнение должностных "
                "обязанностей, предусмотренных настоящей должностной "
                "инструкцией – в пределах, определенных действующим "
                "трудовым законодательством Российской Федерации;",
                "разглашение либо неправомерную обработку персональных данных.",
            ],
        },
        {
            "title": "Заключительные положения",
            "points": [
                "Ознакомление работника с должностной инструкцией "
                "осуществляется при приеме на работу.",
                "Должностная инструкция действует бессрочно.",
            ],
        },
    ]

    path = create_di(
        position_full=(
            "специалиста тестового отдела Федерального государственного "
            "бюджетного учреждения «Научный центр»"
            ""
            ""
        ),
        position_full_genitive="специалиста тестового отдела",
        sections=sample_sections,
        output_path="/tmp/ДИ smoke-test.docx",
    )
    print(f"Тест: {path}")
