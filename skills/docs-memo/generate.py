#!/usr/bin/env python3
"""Генератор служебных записок.

Конфигурация организации загружается из references/org_details.md.
При отсутствии файла запустите docs-init для первичной настройки.

Использование:
    python3 generate.py  # создает служебная_записка.docx с примером
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import re as _re
import warnings as _warnings

_SLASH_OK = _re.compile(r'п/п|ИНН/КПП|[0-9]/[0-9]|[A-Za-z]/[A-Za-z]')

def _warn_slash(text):
    """Предупреждение при обнаружении косой черты между кириллическими словами."""
    for m in _re.finditer(r'[а-яА-ЯёЁ]+\s*/\s*[а-яА-ЯёЁ]+', text):
        if not _SLASH_OK.search(m.group()):
            _warnings.warn(
                f"Косая черта в тексте документа: \u00ab{m.group()}\u00bb. "
                f"Используйте запятую, \u00abили\u00bb, \u00abи\u00bb или скобки.",
                stacklevel=3,
            )

import sys as _sys

_HERE = os.path.dirname(os.path.abspath(__file__))
for _lib_path in (
    os.path.join(_HERE, "..", "..", "lib"),
    os.path.expanduser("~/.docs-plugin/runtime/lib"),
):
    if os.path.isdir(_lib_path) and _lib_path not in _sys.path:
        _sys.path.insert(0, _lib_path)
from db import log_generation, ORG_DETAILS_PATH


def _load_org_details():
    """Загрузить реквизиты из ~/.docs-plugin/org_details.md."""
    config = {}
    if not os.path.exists(ORG_DETAILS_PATH):
        print("WARNING: ~/.docs-plugin/org_details.md not found. Run docs-init to configure.")
        return config
    with open(ORG_DETAILS_PATH, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if ": " in line and not line.startswith("#"):
                key, _, value = line.partition(": ")
                config[key.strip()] = value.strip()
    return config


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_para(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    _warn_slash(text)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    # Совместимость с Word (восточные шрифты)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def create_sluzhebka(
    addressee=None,
    sender=None,
    body_paragraphs=None,
    signer=None,
    doc_date=None,
    output_path="служебная_записка.docx",
):
    """
    Создает служебную записку в формате .docx.

    Параметры
    ----------
    addressee : str | None
        Полный блок «кому» (должность + ФИО в дательном падеже).
        По умолчанию строится из org_details.md.
    sender : str | None
        Блок «от кого» (должность + ФИО в родительном падеже).
        По умолчанию строится из org_details.md.
    body_paragraphs : list[str]
        Список абзацев основного текста.
    signer : str | None
        ФИО подписанта (именительный падеж). По умолчанию из org_details.md.
    doc_date : str | None
        Дата документа в формате «ДД.ММ.ГГГГ». По умолчанию - сегодня.
    output_path : str
        Путь для сохранения .docx.
    """
    o = _load_org_details()

    if addressee is None:
        leader_title = o.get("leader_title", "")
        short_name = o.get("short_name", "")
        leader_name_gen = o.get("leader_name_gen", "")
        addressee = "\n".join(filter(None, [leader_title, short_name, leader_name_gen]))

    if sender is None:
        position_gen = o.get("author_position_gen") or o.get("author_position", "")
        name_short = o.get("author_name_short", "")
        sender = "от " + position_gen + ("\n" + name_short if name_short else "")

    if signer is None:
        signer = o.get("author_name_short", "")

    if body_paragraphs is None:
        body_paragraphs = ["Текст служебной записки."]
    if doc_date is None:
        doc_date = date.today().strftime("%d.%m.%Y")

    doc = Document()

    # Поля страницы (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    # Шрифт Normal по умолчанию
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(14)

    # ── Шапка ──────────────────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=2)
    _remove_table_borders(hdr)
    hdr.columns[0].width = Cm(8)
    hdr.columns[1].width = Cm(8.5)

    _cell_para(hdr.cell(0, 1), addressee)
    _cell_para(hdr.cell(1, 1), sender)

    doc.add_paragraph()  # отступ

    # ── Заголовок ──────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Служебная записка")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_run.bold = True

    doc.add_paragraph()  # отступ

    # ── Основной текст ─────────────────────────────────────────────────────────
    for text in body_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    doc.add_paragraph()  # отступ перед подписью

    # ── Подпись ────────────────────────────────────────────────────────────────
    sig = doc.add_table(rows=1, cols=2)
    _remove_table_borders(sig)
    sig.columns[0].width = Cm(8)
    sig.columns[1].width = Cm(8.5)

    _cell_para(
        sig.cell(0, 1),
        f"{signer}\n\n{doc_date}",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    output_path = os.path.expanduser(output_path)
    doc.save(output_path)
    print(f"Сохранено: {output_path}")
    log_generation("memo", "Служебная записка", output_path, {
        "paragraphs_count": len(body_paragraphs),
    })
    return output_path


if __name__ == "__main__":
    create_sluzhebka(
        body_paragraphs=[
            "Прошу Вас рассмотреть возможность...",
        ],
        doc_date=date.today().strftime("%d.%m.%Y"),
        output_path="служебная_записка_пример.docx",
    )
